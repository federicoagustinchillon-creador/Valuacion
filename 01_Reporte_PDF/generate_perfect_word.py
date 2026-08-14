# -*- coding: utf-8 -*-
"""
generate_perfect_word.py
Genera un documento Word (.docx) institucional de 17 secciones que refleja 1:1 el reporte LaTeX.
Aplica tipografía Georgia, paleta Navy (#0D233A) y Blue (#1E3A8A), tablas estilo booktabs,
cajas sombreadas con borde izquierdo estilo tcolorbox, encabezados y pies de página formales,
todas las 31 figuras perfectamente emparejadas con sus epígrafes, Índice General, Apéndice Contable
completo (EEFF Auditados FY20-FY25, Reconciliación Impositiva, FCFF 2026E-2030E), 13 Referencias
Bibliográficas y Disclaimer Legal/Certificación del Analista.
"""

import os, sys, re, shutil
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(DIR)
FIG_DIR = os.path.join(ROOT_DIR, "03_Modelo_y_Codigo", "figuras")
OUTPUT_DOCX = os.path.join(DIR, "Informe_Valuacion_Aluar_word.docx")
DOWNLOADS_DOCX = r"c:\Users\fedea\Downloads\Informe_Valuacion_Aluar.docx"

HEX_NAVY = "0D233A"
HEX_BLUE = "1E3A8A"
HEX_GRAY = "4A4A4A"
HEX_LIGHTGRAY = "F8FAFC"
HEX_BORDER = "D1D5DB"
HEX_GREEN = "00843D"
HEX_RED = "B91C1C"

COLOR_NAVY = RGBColor(13, 35, 58)
COLOR_BLUE = RGBColor(30, 58, 138)
COLOR_GRAY = RGBColor(74, 74, 74)
COLOR_GREEN = RGBColor(0, 132, 61)
COLOR_RED = RGBColor(185, 28, 28)
COLOR_BLACK = RGBColor(15, 23, 42)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table, color_hex=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.text = ""
        r_left = hp.add_run("ALUAR S.A.I.C. (ALUA.BA) · Equity Research                                                 ")
        r_left.font.name = "Georgia"
        r_left.font.size = Pt(8.5)
        r_left.font.bold = True
        r_left.font.color.rgb = COLOR_NAVY
        
        r_right = hp.add_run("FCE · UNCuyo")
        r_right.font.name = "Georgia"
        r_right.font.size = Pt(8.5)
        r_right.font.color.rgb = COLOR_GRAY

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.text = ""
        r_f1 = fp.add_run("Analista Principal: Federico Agustín Chillón | federico.chillon@fce.uncu.edu.ar")
        r_f1.font.name = "Georgia"
        r_f1.font.size = Pt(8.0)
        r_f1.font.color.rgb = COLOR_GRAY

def create_callout_box(doc, text_list, title=None, border_color=HEX_NAVY, bg_color=HEX_LIGHTGRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(7.0)
    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        rt = p.add_run(f"{title}\n")
        rt.font.name = "Georgia"
        rt.font.size = Pt(10.0)
        rt.font.bold = True
        rt.font.color.rgb = COLOR_NAVY
        
    for item in text_list:
        rb = p.add_run(f"{item}\n")
        rb.font.name = "Georgia"
        rb.font.size = Pt(9.0)
        rb.font.color.rgb = COLOR_BLACK
            
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(13.0)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY
    
    pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="2" w:color="{HEX_NAVY}"/></w:pBrd>')
    p._p.get_or_add_pPr().append(pBrd)

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(11.0)
    r.font.bold = True
    r.font.color.rgb = COLOR_BLUE

def add_body_p(doc, text, bold_prefix=None, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name = "Georgia"
        rb.font.size = Pt(9.5)
        rb.font.bold = True
        rb.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_equation_box(doc, eq_text, eq_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r = p.add_run(f"   {eq_text}   ")
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY
    
    if eq_label:
        rl = p.add_run(f"    ({eq_label})")
        rl.font.name = "Georgia"
        rl.font.size = Pt(9.0)
        rl.font.italic = True
        rl.font.color.rgb = COLOR_GRAY

def add_paired_figures(doc, fig_a_num, fig_a_cap, fig_b_num, fig_b_cap):
    fig_a_path = os.path.join(FIG_DIR, f"figura_{fig_a_num:02d}.png")
    fig_b_path = os.path.join(FIG_DIR, f"figura_{fig_b_num:02d}.png")
    
    if not (os.path.exists(fig_a_path) and os.path.exists(fig_b_path)):
        return
        
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.4)
            set_cell_margins(cell, top=40, bottom=40, left=40, right=40)
            
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(2)
    p0.add_run().add_picture(fig_a_path, width=Inches(3.3))
    
    p1 = table.cell(0, 1).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run().add_picture(fig_b_path, width=Inches(3.3))
    
    c0 = table.cell(1, 0).paragraphs[0]
    c0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c0.paragraph_format.space_after = Pt(4)
    r0 = c0.add_run(f"Figura {fig_a_num}: {fig_a_cap}")
    r0.font.name = "Georgia"
    r0.font.size = Pt(8.0)
    r0.font.color.rgb = COLOR_GRAY
    
    c1 = table.cell(1, 1).paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c1.paragraph_format.space_after = Pt(4)
    r1 = c1.add_run(f"Figura {fig_b_num}: {fig_b_cap}")
    r1.font.name = "Georgia"
    r1.font.size = Pt(8.0)
    r1.font.color.rgb = COLOR_GRAY
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_custom_table(doc, headers, data, caption=None, col_widths=None):
    if caption:
        pc = doc.add_paragraph()
        pc.paragraph_format.space_before = Pt(6)
        pc.paragraph_format.space_after = Pt(2)
        pc.paragraph_format.keep_with_next = True
        rc = pc.add_run(f"Tabla: {caption}")
        rc.font.name = "Georgia"
        rc.font.size = Pt(9.0)
        rc.font.bold = True
        rc.font.color.rgb = COLOR_NAVY
        
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(str(h))
        r.font.name = "Georgia"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY
        set_cell_shading(hdr_cells[i], "F1F5F9")
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = Inches(col_widths[i])
            
    for r_idx, row in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(val))
            r.font.name = "Georgia"
            r.font.size = Pt(8.5)
            if "Total" in str(row[0]) or "ROIC" in str(row[0]) or "EVA" in str(row[0]) or "Target" in str(row[0]) or "WACC" in str(row[0]):
                r.font.bold = True
                r.font.color.rgb = COLOR_NAVY
            else:
                r.font.color.rgb = COLOR_BLACK
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=100, right=100)
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = Inches(col_widths[c_idx])
                
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_word_report():
    print("Construyendo documento Word estructurado en 17 secciones...")
    doc = Document()
    add_header_footer(doc)
    
    # ═════════════════════════════════════════════════════════════════════════
    # HOJA 1: PORTADA / RESUMEN EJECUTIVO
    # ═════════════════════════════════════════════════════════════════════════
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)
    r_t = p_title.add_run("Aluar Aluminio Argentino")
    r_t.font.name = "Georgia"
    r_t.font.size = Pt(22.0)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_NAVY
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(8)
    r_s = p_sub.add_run("Reporte de Valuación Fundamental & Cuantitativa Institucional  |  ALUA.BA")
    r_s.font.name = "Georgia"
    r_s.font.size = Pt(12.0)
    r_s.font.color.rgb = COLOR_GRAY
    
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(8)
    pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="{HEX_NAVY}"/></w:pBrd>')
    p_hr._p.get_or_add_pPr().append(pBrd)
    
    # Summary Box Table (2 columns: Left text, Right metrics)
    box_table = doc.add_table(rows=1, cols=2)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_table.autofit = False
    
    c_left = box_table.cell(0, 0)
    c_right = box_table.cell(0, 1)
    c_left.width = Inches(4.2)
    c_right.width = Inches(2.8)
    set_cell_margins(c_left, top=100, bottom=100, left=120, right=120)
    set_cell_margins(c_right, top=100, bottom=100, left=120, right=120)
    set_cell_shading(c_right, "F8FAFC")
    
    pl = c_left.paragraphs[0]
    pl.paragraph_format.line_spacing = 1.15
    rl_head = pl.add_run("Resumen Ejecutivo & Tesis de Inversión\n")
    rl_head.font.name = "Georgia"
    rl_head.font.size = Pt(11.0)
    rl_head.font.bold = True
    rl_head.font.color.rgb = COLOR_NAVY
    
    rl_body = pl.add_run(
        "El dictamen técnico del modelo cuantitativo para Aluar Aluminio Argentino S.A.I.C. (ALUA.BA) concluye una recomendación de COMPRAR con un precio objetivo fundamental a 12 meses de ARS 1.236,00 (Target Integrado ARS 1.255,60 incorporando la opción real eólica PEAL V), lo que representa un retorno esperado de +25,8% / +27,8% frente a la cotización spot de ARS 982,50.\n\n"
        "1. Liderazgo Estructural en Costos: Matriz 96% autogenerada (78% renovable), blindando su costo en efectivo C1 en USD 1.680/t (primer cuartil global).\n"
        "2. Flujo Dolarizado y Cobertura Natural: 80% de ingresos en USD vía LME frente a costos fijos domésticos en pesos.\n"
        "3. Compresión Soberana y Desapalancamiento: EMBI+ a 441 pb reduce el WACC a 7,06% (λ=0,20), y el cese del CAPEX eólico recomprime la deuda neta a <1,0x EBITDA hacia FY2028E."
    )
    rl_body.font.name = "Georgia"
    rl_body.font.size = Pt(9.0)
    rl_body.font.color.rgb = COLOR_BLACK
    
    pr = c_right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr.paragraph_format.line_spacing = 1.15
    
    rr_stat = pr.add_run("Dictamen del Modelo Cuantitativo\n")
    rr_stat.font.name = "Georgia"
    rr_stat.font.size = Pt(8.5)
    rr_stat.font.color.rgb = COLOR_GRAY
    
    rr_rec = pr.add_run("COMPRAR\n")
    rr_rec.font.name = "Georgia"
    rr_rec.font.size = Pt(16.0)
    rr_rec.font.bold = True
    rr_rec.font.color.rgb = COLOR_GREEN
    
    rr_data = pr.add_run(
        "Target Base (Dumrauf): ARS 1.236,00\n"
        "Opción Real PEAL V: +ARS 19,60\n"
        "Target Integrado: ARS 1.255,60 (USD 0,79)\n"
        "Cotización Spot: ARS 982,50\n"
        "Upside Integrado: +27,8%\n"
        "─────────────────────────\n"
        "WACC Modelo: 7,06%\n"
        "Costo Equity (Ke): 9,30%\n"
        "Beta Apalancado: 0,888\n"
        "ERP Damodaran: 4,18%\n"
        "CRP (λ × EMBI+): 0,88%\n"
        "Kd (pre / post-tax): 3,80% / 2,47%\n"
        "D/E Objetivo: 0,486\n"
        "Crecimiento Terminal (g): 2,00%\n"
        "Tipo de Cambio CCL: 1.584,25"
    )
    rr_data.font.name = "Georgia"
    rr_data.font.size = Pt(8.5)
    rr_data.font.color.rgb = COLOR_NAVY
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    create_callout_box(
        doc,
        ["Documento de investigación cuantitativa y análisis financiero con fines exclusivamente educativos y de divulgación académica. No constituye una recomendación de compra, venta o asesoramiento de inversión en los términos de la Ley N.° 26.831. Las opiniones técnicas representan estimaciones independientes fundamentadas en modelos estocásticos y estados financieros públicos auditados."],
        title="Aviso Legal y Descargo de Responsabilidad"
    )
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════════
    # HOJA 2: ÍNDICE GENERAL
    # ═════════════════════════════════════════════
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_th = p_toc_head.add_run("Índice General de Secciones y Contenidos")
    r_th.font.name = "Georgia"
    r_th.font.size = Pt(14.0)
    r_th.font.bold = True
    r_th.font.color.rgb = COLOR_NAVY
    
    toc_items = [
        ("1. Descripción de la Compañía: Integración Vertical y Estructura Operativa", "3"),
        ("2. Contexto Macroeconómico y Mecánica Operativa de Proyección (P × Q)", "4"),
        ("3. Dinámica de la Industria Global del Aluminio y Cointegración con el LME", "5"),
        ("4. Gobernanza Corporativa, Marco Regulatorio y Posicionamiento ESG", "6"),
        ("5. Valuación Relativa y Dinámica de Desapalancamiento Post-CAPEX", "7"),
        ("6. Costo de Capital (WACC) y Fundamentación Económica del Factor λ", "8"),
        ("7. Análisis de Creación de Valor (EVA) y Disciplina de Convergencia Terminal", "9"),
        ("8. Puente de Valuación: Desglose de Enterprise Value a Equity Value", "10"),
        ("9. Análisis Estratégico DuPont y Frameworks Industriales (Porter / PESTEL)", "11"),
        ("10. Análisis Financiero Histórico, Ratios Operativos y Cobertura", "12"),
        ("11. Valuación Fundamental por Flujo de Fondos Descontados (DCF) y Opción Real", "13"),
        ("12. Análisis de Sensibilidad Bidimensional y Matriz de Rangos", "14"),
        ("13. Simulación Estocástica de Monte Carlo y DCF Inverso", "14"),
        ("14. Valuación Relativa Sectorial y Múltiplos Comparables LTM y Forward", "15"),
        ("15. Gestión Cuantitativa de Riesgo y Asignación de Portafolio", "15"),
        ("16. Modelización Estocástica, Dependencia de Cola y Validación Econométrica", "16"),
        ("17. Dictamen Final y Recomendación de Inversión", "17"),
        ("Apéndice: Estados Financieros Auditados y Proyecciones de Flujo de Fondos (USD MM)", "18"),
        ("Referencias Bibliográficas y Fuentes Académicas", "19"),
        ("Aviso Legal y Certificación del Analista (Research Standards / Ley 26.831)", "19")
    ]
    
    toc_table = doc.add_table(rows=len(toc_items), cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    toc_table.autofit = False
    for i, (title_text, page_num) in enumerate(toc_items):
        r_c0 = toc_table.cell(i, 0)
        r_c1 = toc_table.cell(i, 1)
        r_c0.width = Inches(6.0)
        r_c1.width = Inches(1.0)
        set_cell_margins(r_c0, top=30, bottom=30, left=40, right=40)
        set_cell_margins(r_c1, top=30, bottom=30, left=40, right=40)
        
        p0 = r_c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(title_text)
        r0.font.name = "Georgia"
        r0.font.size = Pt(8.5)
        if title_text.startswith("Apéndice") or title_text.startswith("Referencias") or title_text.startswith("Aviso"):
            r0.font.bold = True
            r0.font.color.rgb = COLOR_BLUE
        elif any(title_text.startswith(f"{n}.") for n in range(1, 18)):
            r0.font.bold = True
            r0.font.color.rgb = COLOR_NAVY
            
        p1 = r_c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(page_num)
        r1.font.name = "Georgia"
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = COLOR_GRAY
        
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 1: DESCRIPCIÓN DE LA COMPAÑÍA
    # ═════════════════════════════════════════════
    add_heading_1(doc, "1. Descripción de la Compañía: Integración Vertical y Estructura Operativa")
    add_heading_2(doc, "1.1 Perfil Operativo: Integración Vertical y Capacidad al 100%")
    add_body_p(doc, "Aluar Aluminio Argentino S.A.I.C. es el único productor primario de aluminio integrado en Argentina y uno de los mayores complejos industriales de Sudamérica. Fundada en 1974, la compañía opera su planta de reducción electrolítica en Puerto Madryn (Provincia del Chubut) con una capacidad instalada nominal de 460.000 toneladas anuales. Su portafolio abarca aluminio primario líquido y sólido (lingotes estándar, lingotes T, barrotes de extrusión, alambrón y placas de laminación), junto con una división de semielaborados y elaborados de alto valor agregado. Alrededor del 80% del volumen físico producido se exporta a mercados internacionales (Estados Unidos, Brasil, Europa y Japón), liquidado en dólares estadounidenses sobre cotizaciones del London Metal Exchange (LME) más primas regionales, lo que dota a la firma de un flujo operativo estructuralmente dolarizado.")
    
    add_paired_figures(
        doc,
        1, "Hitos Históricos de Aluar (1974–2026): Trayectoria de expansión de capacidad instalada y transición energética.",
        2, "Estructura Accionaria: Concentración del 45,98% en el grupo de control (Familia Madanes) y 54,02% capital flotante/ANSES."
    )
    
    add_heading_2(doc, "1.2 Estructura Accionaria y Control Corporativo")
    add_body_p(doc, "Según la Memoria Anual y Estados Financieros al 30-jun-2025, el capital social está integrado por 2.800 millones de acciones ordinarias escriturales de $1 valor nominal y un voto por acción, listadas en BYMA bajo la especie ALUA.BA. Las sociedades vinculadas a la familia Madanes (CIPAL S.A., Estancias San Javier de Catamarca S.A. y Angerona Group Trust) concentran en conjunto el 45,98% de los derechos de voto, conformando una mayoría de control efectiva frente al 54,02% atomizado entre la ANSES-FGS (27,47%) y el capital flotante en el mercado doméstico e internacional.")
    
    add_heading_2(doc, "1.3 Ventajas Competitivas: Matriz Energética y Posición de Costos C1")
    add_body_p(doc, "La fundición de aluminio primario es un proceso electrointensivo donde la energía representa entre el 30% y 40% del costo total de producción. Tres pilares fundamentan el foso defensivo de Aluar:")
    add_body_p(doc, "1. Autonomía Energética Propia (96% del Consumo): La planta de Puerto Madryn es abastecida en un 54% por la Central Hidroeléctrica Futaleufú (320 MW asignados contractualmente bajo concesión), 24% por el Parque Eólico Aluar (PEAL, 200 MW en operación) y 18% por ciclos combinados térmicos propios a gas natural. Sólo el 4% restante se toma de la red nacional (SADI/CAMMESA), blindando a la compañía ante aumentos en las tarifas eléctricas spot.")
    add_body_p(doc, "2. Posición en el Primer Cuartil de Costos Globales: Con un costo en efectivo C1 de USD 1.680 por tonelada (percentil 29% de la curva global de costos CRU/Wood Mackenzie), Aluar se ubica entre los productores más eficientes del mundo, asegurando márgenes operativos positivos aún en los valles del ciclo del LME.")
    add_body_p(doc, "3. Infraestructura Portuaria Exclusiva: Terminal portuaria propia de aguas profundas en Puerto Madryn para la descarga directa de bauxita/alúmina importada y despacho de exportaciones con economías de escala en fletes marítimos.")
    
    add_heading_2(doc, "1.4 Estrategia Corporativa: Certificación ASI y Mezcla de Valor Agregado")
    add_body_p(doc, "Los objetivos estratégicos comprenden: (1) alcanzar el 100% de abastecimiento renovable mediante la Etapa V del Parque Eólico (PEAL V / La Flecha); (2) obtener la certificación ASI (Aluminium Stewardship Initiative) para capturar primas de precio por aluminio verde en la Unión Europea; y (3) defender la cuota en productos elaborados:")
    
    add_custom_table(
        doc,
        ["Métrica de Volumen", "FY2023", "FY2024", "FY2025"],
        [
            ["Ventas Elaborados (t)", "23.104", "17.766", "13.498"],
            ["Ventas Totales (t)", "383.834", "418.014", "402.991"],
            ["Participación Elaborados", "6,0%", "4,3%", "3,3%"]
        ],
        caption="Participación de la División Elaborados en el Volumen Total de Ventas (Memoria y EEFF)",
        col_widths=[3.0, 1.3, 1.3, 1.3]
    )
    
    add_heading_2(doc, "1.5 Análisis Estratégico FODA")
    add_body_p(doc, "• Fortalezas: Monopolio primario integrado en Argentina; 96% de energía autogenerada; costos C1 en primer cuartil mundial; ingresos 80% dolarizados.")
    add_body_p(doc, "• Oportunidades: Captura de primas comerciales por certificación ASI; incentivos impositivos del RIGI para PEAL V; compresión adicional del riesgo soberano.")
    add_body_p(doc, "• Debilidades: Dependencia total de alúmina importada; capacidad nominal saturada al 100% (460 kt); apalancamiento transitorio por inversiones eólicas.")
    add_body_p(doc, "• Amenazas: Volatilidad cíclica del LME; apreciación del tipo de cambio real que encarece costos laborales en USD; sobrecapacidad de fundición en Asia.")
    
    add_heading_2(doc, "1.6 Desempeño Financiero Reciente (FY2020–FY2025)")
    add_custom_table(
        doc,
        ["Métrica Financiera", "FY2020", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025"],
        [
            ["Ingresos (Ventas Netas)", "823,5", "513,5", "654,7", "647,8", "915,9", "1.092,6"],
            ["EBITDA", "116,5", "107,4", "209,7", "105,7", "204,7", "163,3"],
            ["Margen EBITDA", "14,2%", "20,9%", "32,0%", "16,3%", "22,4%", "14,9%"],
            ["Deuda Neta", "351,3", "137,4", "57,3", "169,9", "211,2", "525,8"],
            ["Deuda Neta / EBITDA", "3,02x", "1,28x", "0,27x", "1,61x", "1,03x", "3,22x"]
        ],
        caption="Evolución de Indicadores Clave Auditados FY2020–FY2025 (USD MM)",
        col_widths=[2.4, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 2: CONTEXTO MACROECONÓMICO Y MECÁNICA OPERATIVA (P x Q)
    # ═════════════════════════════════════════════
    add_heading_1(doc, "2. Contexto Macroeconómico y Mecánica Operativa de Proyección (P × Q)")
    add_heading_2(doc, "2.1 Entorno Macroeconómico y Dinámica del Riesgo Soberano")
    add_body_p(doc, "Las proyecciones del modelo se apoyan en el marco macroeconómico de estabilización argentino, con convergencia inflacionaria y recuperación del PBI. La compresión del spread de riesgo soberano EMBI+ desde máximos de 2.400 pb en 2022 a 441 pb al cierre de julio de 2026 ejerce un impacto directo y favorable sobre la tasa de descuento del capital.")
    
    add_paired_figures(
        doc,
        3, "Contexto Macroeconómico: PBI real e inflación proyectada en Argentina (2020–2030E).",
        4, "Compresión del Riesgo País (EMBI+): Descenso estructural del spread soberano desde 2.400 pb a 441 pb."
    )
    
    add_heading_2(doc, "2.2 Mecánica Operativa de Ingresos: Volumen Saturado y Reversión del LME")
    add_body_p(doc, "La proyección de ingresos descansa en la formulación microeconómica Ingresos = P × Q:")
    add_equation_box(doc, "Q = 460.000 Tn × 0,94 = 432.400 Tn/año", "Volumen Físico")
    add_body_p(doc, "En FY2026E los ingresos alcanzan USD 1.591,4 MM escalados por el nivel spot del LME en el 1T26 ($3.173,2/t). Para el horizonte explícito 2027E–2030E, el precio realizado converge linealmente a su valor de régimen de equilibrio de largo plazo:")
    add_equation_box(doc, "P_2030 = LME_rev ($2.587,1/t) + Premium_elab ($645,7/t) = USD 3.232,8/t", "Precio 2030E")
    
    add_body_p(doc, "Se calibra el parámetro de eficiencia k sobre la observación real del 1T26 (Ventas = USD 416 MM, EBITDA = USD 106 MM, C1 = USD 1.680/t):")
    add_equation_box(doc, "k = Margen_EBITDA_Q1 / ((P_Q1 - C1) / P_Q1) = (106 / 416) / ((3.848,3 - 1.680) / 3.848,3) = 0,4521", "Calibración k")
    
    add_body_p(doc, "La variación de capital de trabajo se modela en base a la mediana histórica (NWC/Ventas = 49,42%):")
    add_equation_box(doc, "ΔNWC_2026E = (Ventas_2026E - Ventas_2025) × 49,42% = (1.591,4 - 1.092,6) × 0,4942 = USD 246,5 MM", "ΔNWC 2026E")
    
    add_heading_2(doc, "2.3 Curva de Tasas y Convergencia del Riesgo País (Modelo AR(1))")
    add_equation_box(doc, "EMBI_t = μ × (1 - φ) + φ × EMBI_{t-1} + ε_t,   φ = 0,7582,   μ = 1.614 pb", "AR(1) Riesgo País")
    
    add_paired_figures(
        doc,
        5, "Curva de Riesgo Soberano: Convergencia estocástica AR(1) frente al supuesto corporativo base (441 pb).",
        6, "Aluminio LME vs. DXY: Desacoplamiento cíclico y correlación negativa frente al índice del dólar global."
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 3: DINÁMICA DE LA INDUSTRIA GLOBAL
    # ═════════════════════════════════════════════
    add_heading_1(doc, "3. Dinámica de la Industria Global del Aluminio y Cointegración con el LME")
    add_heading_2(doc, "3.1 Desacoplamiento Cíclico: LME, DXY y Cobertura Cambiaria Natural")
    add_body_p(doc, "Existe una correlación inversa estructural entre la cotización global de los commodities y el índice del dólar (DXY). Aluar cuenta con una cobertura natural: cuando el ciclo internacional del commodity se deprime y el dólar se fortalece, los episodios de devaluación del tipo de cambio real en Argentina licúan los costos operativos fijos en pesos (mano de obra y servicios locales), amortiguando la compresión de márgenes.")
    
    add_heading_2(doc, "3.2 Diagnóstico Econométrico: Test de Cointegración de Engle-Granger")
    add_body_p(doc, "Se sometió la serie diaria de ALUA.BA (homogeneizada a USD vía CCL, N=2.355 ruedas, 2016–2026) y el futuro del LME a prueba de cointegración formal de Engle-Granger. El estadístico t = -2,29 (p = 0,38, valor crítico al 5% de -3,34) rechaza la cointegración estricta en niveles. No obstante, la elasticidad contemporánea es de 0,66 (R² = 0,12). Esto confirma que Aluar no es un activo puramente ligado al commodity ni puramente dependiente del riesgo argentino, sino un activo dual con dinámica mixta.")
    
    add_paired_figures(
        doc,
        7, "Producción Global de Aluminio: Participación de los principales países productores mundiales.",
        8, "Escala y Autogeneración: Comparativa de capacidad y abastecimiento energético propio frente a gigantes."
    )
    
    add_paired_figures(
        doc,
        9, "Cuota de Mercado Doméstica: Participación del 60% en el mercado argentino frente a importaciones.",
        10, "Matriz Energética y Curva C1: Abastecimiento renovable y posición defensiva en el primer cuartil."
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 4: GOBERNANZA CORPORATIVA Y ESG
    # ═════════════════════════════════════════════
    add_heading_1(doc, "4. Gobernanza Corporativa, Marco Regulatorio y Posicionamiento ESG")
    add_heading_2(doc, "4.1 Estructura de Gobierno y Control Interno (Ley 26.831 / CNV)")
    add_body_p(doc, "Aluar opera bajo el régimen de oferta pública de la CNV (Resolución General 797/2019). El Directorio está presidido por Javier Madanes Quintanilla. El Comité de Auditoría cuenta con mayoría de miembros independientes conforme al Art. 109 de la Ley 26.831, dictaminando sobre el control interno y transacciones con partes vinculadas como Hidroeléctrica Futaleufú S.A. (60,2%) y Transpa S.A. (20,4%), verificadas a precios y condiciones de mercado.")
    
    add_custom_table(
        doc,
        ["Dimensión", "Métricas y Evidencia Observable", "Evaluación Cualitativa"],
        [
            ["Ambiental (E)", "78% renovable autogenerada. Huella < 4 t CO2/t Al. Proceso de certificación ASI.", "Fuerte. Ventaja competitiva frente al arancel fronterizo de carbono europeo (CBAM)."],
            ["Social (S)", "Principal empleador privado de Chubut (> 6.200 empleos directos e indirectos). TRIR < 1,0.", "Adecuada. Riguroso estándar de seguridad industrial y baja conflictividad laboral."],
            ["Gobernanza (G)", "Concentración del 45,98% en el grupo de control. Auditoría externa por PwC e independencia de comités.", "Moderada. Estabilidad estratégica de largo plazo con flotante acotado al 54,02%."]
        ],
        caption="Scorecard Proxy de ESG y Gobernanza Corporativa",
        col_widths=[1.5, 3.2, 2.3]
    )
    
    add_heading_2(doc, "4.2 Mitigación del Impuesto Fronterizo por Carbono (CBAM de la UE)")
    add_body_p(doc, "La intensidad de carbono de Aluar (< 4 t CO2/t Al) se ubica muy por debajo del promedio global (~ 11 t CO2/t), dominado por fundiciones chinas e indias alimentadas a carbón. Esto le otorga una ventaja arancelaria neta estimada en USD 80–120 por tonelada bajo el esquema CBAM de la Unión Europea.")
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 5: VALUACIÓN RELATIVA Y DESAPALANCAMIENTO
    # ═════════════════════════════════════════════
    add_heading_1(doc, "5. Valuación Relativa y Dinámica de Desapalancamiento Post-CAPEX")
    add_heading_2(doc, "5.1 Múltiplos LTM vs. Forward: Compresión Post-Recuperación Operativa")
    add_body_p(doc, "Aluar transa a un múltiplo EV/EBITDA LTM de 13,4x sobre el EBITDA deprimido de FY2025 (USD 163,3 MM), reflejando una prima aparente frente a comparables globales (mediana de 6,5x). Sin embargo, al valuar sobre el EBITDA proyectado normalizado de FY2026E (USD 391,2 MM), el múltiplo implícito de Aluar colapsa a 5,2x EV/EBITDA Forward, situando a la compañía con un descuento fundamental del 20%–25% respecto de pares como Alcoa (8,1x) y Norsk Hydro (6,5x).")
    
    add_paired_figures(
        doc,
        11, "Múltiplos EV/EBITDA vs. Margen: Posicionamiento de Aluar frente a pares internacionales (Alcoa, Hydro, Chalco).",
        12, "EBITDA Histórico y Proyectado: Recuperación operativa y desapalancamiento post-pico de CAPEX PEAL V."
    )
    
    add_heading_2(doc, "5.2 Desapalancamiento Post-Pico de Inversión Eólica")
    add_body_p(doc, "El CAPEX de FY2025 (USD 243,9 MM) elevó la Deuda Neta a USD 525,8 MM (3,22x EBITDA). Con el cese del desembolso intensivo y la estabilización del CAPEX de mantenimiento en ~USD 90–103 MM/año, la generación de caja proyectada comprime el ratio Deuda Neta / EBITDA a <1,0x hacia FY2028E.")
    
    add_custom_table(
        doc,
        ["Compañía", "Ticker", "EV/EBITDA (x)", "Margen EBITDA", "Deuda Neta/EBITDA"],
        [
            ["Aluar S.A.I.C. (LTM / Fwd)", "ALUA.BA", "13,4x / 5,2x", "14,9%", "3,22x"],
            ["Alcoa Corp.", "AA", "8,1x", "9,2%", "2,10x"],
            ["Norsk Hydro", "NHY.OL", "6,5x", "11,4%", "1,40x"],
            ["Chalco", "2600.HK", "5,8x", "8,8%", "3,80x"],
            ["Mediana Industria", "", "6,5x", "9,2%", "2,10x"]
        ],
        caption="Análisis de Comparables Globales de la Industria del Aluminio (julio-2026)",
        col_widths=[2.5, 1.2, 1.2, 1.1, 1.0]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 6: MODELO WACC Y FACTOR LAMBDA
    # ═════════════════════════════════════════════
    add_heading_1(doc, "6. Costo de Capital (WACC) y Fundamentación Económica del Factor λ")
    add_heading_2(doc, "6.1 Formulación Completa del WACC y Parámetros del Modelo")
    add_equation_box(doc, "WACC = Ke × (E / V) + Kd × (1 - t) × (D / V)", "WACC Canónico")
    add_equation_box(doc, "WACC = 9,30% × 0,6729 + 3,80% × (1 - 0,35) × 0,3271 = 6,26% + 0,81% = 7,06%", "Sustitución WACC")
    
    add_heading_2(doc, "6.2 Secuencia del Beta en Cuatro Pasos (OLS → Blume → Hamada)")
    add_body_p(doc, "1. Beta OLS Histórico: Regresión lineal diaria de 10 años (N=2.376 ruedas) de ALUA (USD vía CCL) contra el S&P 500 arrojó β_OLS = 0,8420 (SE = 0,0563, R² = 8,60%).")
    add_body_p(doc, "2. Ajuste de Blume (1975): Corrección por tendencia a la media del riesgo sistemático: β_Blume = (2/3) × 0,8420 + 1/3 = 0,8947.")
    add_body_p(doc, "3. Desapalancamiento por Hamada (1972): Con ratio D/E histórico mediano (0,5036): β_U = 0,8947 / [1 + (1 - 0,35) × 0,5036] = 0,6745.")
    add_body_p(doc, "4. Reapalancamiento al D/E Objetivo (0,4860): β_L = 0,6745 × [1 + (1 - 0,35) × 0,4860] = 0,8876.")
    
    add_paired_figures(
        doc,
        13, "Secuencia del Beta en 4 Pasos: OLS (0,842) → Blume (0,895) → Desapalancado (0,675) → Hamada (β_L = 0,8876).",
        14, "Descomposición WACC Canónico: Ke = 9,30%, Kd = 2,47%, resultando en un WACC de 7,06% en USD."
    )
    
    add_heading_2(doc, "6.3 Fundamentación Económica del Factor λ = 0,20")
    add_equation_box(doc, "Ke = Rf + β_L × ERP + λ × EMBI+ = 4,70% + 0,8876 × 4,18% + 0,20 × 4,41% = 9,30%", "CAPM con λ")
    add_body_p(doc, "El factor λ = 0,20 coincide con la proporción de ingresos destinados al mercado interno (20%). Dado que el 80% restante se exporta a precios LME y los costos operativos fijos están denominados en pesos (mano de obra y fletes), una devaluación del peso licúa los costos en dólares de Aluar, otorgando una cobertura natural contra el riesgo soberano.")
    
    add_custom_table(
        doc,
        ["Parámetro de Riesgo País", "λ = 0,20 (Base)", "λ = 0,50", "λ = 0,75", "λ = 1,00 (Tradicional)"],
        [
            ["Prima por Riesgo País (λ × EMBI+)", "0,88%", "2,21%", "3,31%", "4,41%"],
            ["Costo del Equity (Ke)", "9,30%", "10,63%", "11,73%", "12,83%"],
            ["WACC Resultante (USD)", "7,06%", "7,96%", "8,70%", "9,45%"],
            ["Precio Objetivo Base (ARS)", "1.236,00", "1.042,50", "915,00", "812,00"],
            ["Upside vs. Spot (ARS 982,50)", "+25,8%", "+6,1%", "-6,9%", "-17,4%"]
        ],
        caption="Análisis de Sensibilidad: Factor λ de Exposición al Riesgo País",
        col_widths=[2.6, 1.1, 1.1, 1.1, 1.1]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 7: ANÁLISIS DE CREACIÓN DE VALOR (EVA)
    # ═════════════════════════════════════════════
    add_heading_1(doc, "7. Análisis de Creación de Valor (EVA) y Disciplina de Convergencia Terminal")
    add_heading_2(doc, "7.1 Dinámica del EVA y Efecto del Capital Invertido en Obras en Curso")
    add_body_p(doc, "El Valor Económico Agregado (EVA_t = (ROIC_t - WACC) × NOA_{t-1}) refleja la rentabilidad del capital operativo. En FY2025 el EVA se situó en -USD 60,3 MM debido a la incorporación de USD 243,9 MM en activos eólicos al capital invertido (NOA = USD 1.704,3 MM) antes de que comiencen a generar ahorros operativos plenos.")
    
    add_custom_table(
        doc,
        ["Métrica de Valor", "FY20", "FY21", "FY22", "FY23", "FY24", "FY25", "Terminal (2030E)"],
        [
            ["NOPAT (USD MM)", "31,3", "39,9", "105,4", "40,3", "99,2", "60,0", "135,5"],
            ["Capital Invertido (NOA)", "838,4", "531,3", "590,5", "771,6", "1.248,9", "1.704,3", "1.919,3"],
            ["ROIC (NOPAT / NOA)", "3,73%", "7,50%", "17,85%", "5,22%", "7,94%", "3,52%", "7,06%"],
            ["WACC", "7,06%", "7,06%", "7,06%", "7,06%", "7,06%", "7,06%", "7,06%"],
            ["Spread (ROIC - WACC)", "-3,33%", "+0,44%", "+10,79%", "-1,84%", "+0,88%", "-3,54%", "0,00%"],
            ["EVA (USD MM)", "-27,9", "2,3", "63,7", "-14,2", "11,0", "-60,3", "0,0"]
        ],
        caption="Matriz Histórica y Proyectada de Creación de Valor (EVA, USD MM)",
        col_widths=[2.4, 0.75, 0.75, 0.75, 0.75, 0.75, 0.75, 1.1]
    )
    
    add_heading_2(doc, "7.2 Distinción Metodológica: Caso Base Dumrauf vs. Variante Damodaran")
    add_body_p(doc, "• Modelo Canónico Base (Dumrauf, Cap. 14): En estado estacionario de capacidad saturada (460 kt), CAPEX = D&A y ΔNWC = 0, por lo que FCFF_terminal = NOPAT_2030 = USD 135,5 MM. El valor terminal es TV = (135,5 × 1,02) / (0,0706 - 0,02) = USD 2.731,4 MM, arrojando el target oficial de ARS 1.236,00.")
    add_body_p(doc, "• Variante de Robustez Académica (Damodaran, ROIC → WACC): Si se asume reinversión implícita RR = g / WACC = 2,0% / 7,06% = 28,3%, el flujo terminal se ajusta a NOPAT × (1 - RR) = USD 97,1 MM, resultando en un precio objetivo de ARS 993,00, que se incluye como cota inferior de prudencia.")
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 8: PUENTE DE VALUACIÓN
    # ═════════════════════════════════════════════
    add_heading_1(doc, "8. Puente de Valuación: Desglose de Enterprise Value a Equity Value")
    add_heading_2(doc, "8.1 Estructura del Enterprise Value y Concentración en el Valor Terminal")
    add_body_p(doc, "El Enterprise Value resultante asciende a USD 3.951 MM. El valor presente de los flujos explícitos (2026E–2030E) representa USD 680 MM (17,2% del EV), mientras que el Valor Terminal descontado aporta USD 2.902 MM (73,4% del EV), consistente con la naturaleza intensiva en capital de un activo industrial con vida útil indefinida. Deduciendo la Deuda Financiera Neta auditada de USD 525,8 MM y sumando activos no operativos, se obtiene un Equity Value de USD 2.186 MM, equivalente a ARS 1.236,00 por acción al tipo de cambio CCL de 1.584,25.")
    
    add_paired_figures(
        doc,
        15, "Puente de Valuación (Waterfall): De Enterprise Value (USD 3.951 MM) a Equity Value (USD 2.186 MM) y Target.",
        16, "Días de Capital de Trabajo: Evolución de DIO, DSO, DPO y Ciclo de Conversión de Efectivo (CCC)."
    )
    
    add_heading_2(doc, "8.2 Dinámica del Ciclo de Conversión de Efectivo (CCC)")
    add_body_p(doc, "El Ciclo de Conversión de Efectivo se comprime desde 79 días en FY2023 a 64 días proyectados, impulsado por la reducción en días de inventario (DIO) y optimización en días de cobranza (DSO), liberando capital de trabajo operativo para el servicio de la deuda.")
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 9: ANÁLISIS ESTRATÉGICO DUPONT Y FRAMEWORKS
    # ═════════════════════════════════════════════
    add_heading_1(doc, "9. Análisis Estratégico DuPont y Frameworks Industriales")
    add_heading_2(doc, "9.1 Descomposición DuPont de Cinco Factores (FY2020–FY2025)")
    add_equation_box(doc, "ROE = (RN / EBT) × (EBT / EBIT) × (EBIT / Ventas) × (Ventas / Activos) × (Activos / PN)", "DuPont 5 Factores")
    add_body_p(doc, "En FY2025, la carga impositiva cayó drásticamente a 15,67% (alícuota efectiva del 84,3% por ajuste por inflación impositivo NIC 29), deprimiendo el ROE al 0,83% pese a que el margen operativo EBIT se mantuvo sólido en 8,44%.")
    
    add_custom_table(
        doc,
        ["Factor DuPont", "FY2020", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025"],
        [
            ["Margen Neto (RN / Ventas)", "-5,34%", "5,48%", "15,79%", "19,48%", "9,80%", "0,84%"],
            ["Rotación Activos (V / Activos)", "0,78x", "0,74x", "0,79x", "0,71x", "0,63x", "0,55x"],
            ["Apalancamiento (Activos / PN)", "2,27x", "2,00x", "1,82x", "1,68x", "1,74x", "1,78x"],
            ["ROE (Producto)", "-9,49%", "8,04%", "22,64%", "23,25%", "10,68%", "0,83%"]
        ],
        caption="Descomposición DuPont de 3 Factores (FY2020–FY2025)",
        col_widths=[2.6, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8]
    )
    
    add_heading_2(doc, "9.2 Las Cinco Fuerzas de Porter y Régimen RIGI")
    add_body_p(doc, "• Amenaza de Nuevos Entrantes (Muy Baja): Intensidad de capital requerida (> USD 2.500 MM para un nuevo smelter) y barrera regulatoria insalvable en concesiones hidroeléctricas.")
    add_body_p(doc, "• Poder de Negociación de Proveedores (Bajo a Medio): Dependencia de bauxita/alúmina internacional mitigada por diversificación de orígenes marítimos; energía 96% autogenerada.")
    add_body_p(doc, "• Poder de Negociación de Clientes (Bajo): Mercado externo atomizado con precios tomadores LME; posición monopólica del 60% en el mercado local.")
    add_body_p(doc, "• Amenaza de Sustitutos (Baja): El aluminio conserva ventajas insustituibles en relación peso-resistencia, conductividad y reciclabilidad infinita frente al acero y polímeros.")
    add_body_p(doc, "• Rivalidad de la Industria (Media): Competencia global por costos marginales; blindaje defensivo en el primer cuartil C1.")
    add_body_p(doc, "• Análisis PESTEL / Régimen RIGI: El marco RIGI reduce la tasa del impuesto a las ganancias al 25% para nuevos proyectos eólicos, exime de aranceles de importación a bienes de capital y garantiza estabilidad cambiaria, reduciendo el costo de capital de expansiones futuras.")
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 10: ANÁLISIS FINANCIERO HISTÓRICO Y RATIOS
    # ═════════════════════════════════════════════
    add_heading_1(doc, "10. Análisis Financiero Histórico, Ratios Operativos y Cobertura")
    add_heading_2(doc, "10.1 Evolución Operativa y Solvencia Crediticia")
    add_body_p(doc, "La evolución financiera auditada evidencia la resiliencia de la estructura de costos de Aluar a través de las fluctuaciones del ciclo. La cobertura de intereses (EBIT/Intereses) se mantuvo en un saludable 3,53x en FY2025 incluso en el pico de endeudamiento financiero, con ratios de liquidez corriente superiores a 1,40x:")
    
    add_custom_table(
        doc,
        ["Métrica / Año", "FY2020", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025"],
        [
            ["Ventas Netas", "823,50", "513,50", "654,72", "647,75", "915,85", "1.092,60"],
            ["EBITDA", "116,50", "107,39", "209,67", "105,70", "204,70", "163,25"],
            ["Margen EBITDA", "14,15%", "20,91%", "32,02%", "16,32%", "22,35%", "14,94%"],
            ["EBIT", "48,10", "61,33", "162,14", "61,97", "152,57", "92,27"],
            ["Margen EBIT", "5,84%", "11,94%", "24,76%", "9,57%", "16,66%", "8,44%"],
            ["Resultado Neto", "-43,94", "28,13", "103,35", "126,19", "89,74", "9,17"],
            ["Margen Neto", "-5,34%", "5,48%", "15,79%", "19,48%", "9,80%", "0,84%"],
            ["CAPEX", "-103,01", "-9,43", "-11,11", "-65,89", "-25,59", "-243,93"],
            ["Deuda Neta", "351,30", "137,44", "57,33", "169,88", "211,23", "525,76"],
            ["ROIC", "3,73%", "7,50%", "17,85%", "5,22%", "7,94%", "3,52%"],
            ["ROE", "-9,49%", "8,04%", "22,64%", "23,25%", "10,68%", "0,83%"],
            ["Liquidez Corriente", "1,89x", "3,78x", "2,92x", "2,92x", "4,93x", "2,44x"],
            ["Deuda Neta / EBITDA", "3,02x", "1,28x", "0,27x", "1,61x", "1,03x", "3,22x"],
            ["Cobertura Intereses (EBIT/Int.)", "1,81x", "4,12x", "107,80x", "7,69x", "8,62x", "3,53x"]
        ],
        caption="Histórico Financiero y Ratios Operativos de Aluar S.A.I.C. (USD MM)",
        col_widths=[2.4, 0.75, 0.75, 0.75, 0.75, 0.75, 0.75]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 11: VALUACIÓN FUNDAMENTAL (DCF)
    # ═════════════════════════════════════════════
    add_heading_1(doc, "11. Valuación Fundamental por Flujo de Fondos Descontados (DCF)")
    add_heading_2(doc, "11.1 Estructura del Modelo DCF y Valoración de la Opción Real PEAL V")
    add_equation_box(doc, "EV = Σ [FCFF_{2026+t} / (1 + WACC)^t] + TV / (1 + WACC)^4 = USD 3.951,0 MM", "Descuento DCF")
    add_body_p(doc, "• Precio Objetivo DCF Base (Determinístico): ARS 1.236,00 por acción (USD 0,78, +25,8% Upside).")
    add_body_p(doc, "• Valoración de la Opción Real PEAL V por LSMC: Modela la flexibilidad americana de inversión en 312 MW eólicos mediante regresión por mínimos cuadrados de Longstaff-Schwartz (2001) sobre polinomios de Laguerre, aportando una prima aditiva de +ARS 19,60 por acción (+USD 0,01).")
    add_body_p(doc, "• Precio Objetivo Integrado (Base + Opción Real): ARS 1.255,60 por acción (USD 0,79, +27,8% Upside).")
    
    add_paired_figures(
        doc,
        17, "Test de Estrés Soberano: Impacto sobre WACC y Target en escenario de EMBI+ a 2.400 pb.",
        18, "Football Field de Valuación: Rango de precios intrínsecos por metodología vs. cotización Spot."
    )
    
    add_heading_2(doc, "11.2 Marco de Escenarios Ponderados (Bear / Base / Bull)")
    add_custom_table(
        doc,
        ["Variable / Factor", "Bear (25%)", "Base (50%)", "Bull (25%)"],
        [
            ["Precio LME", "Reversión hacia cash cost global (~USD 1.900/t).", "Media de equilibrio OU (USD 2.450–2.814/t).", "LME sostenido > USD 3.000/t por déficit global."],
            ["Riesgo País (EMBI+)", "Repunte soberano a > 800 pb.", "Convergencia a largo plazo (≈ 441 pb).", "Compresión continua por debajo de 350 pb."],
            ["WACC / g", "9,0% / 1,5%", "7,06% / 2,0%", "6,5% / 2,2%"],
            ["Precio Objetivo (DCF)", "ARS 781 (-20,5%)", "ARS 1.236,00 (Caso Oficial)", "ARS 1.481 (+50,7%)"]
        ],
        caption="Marco de Escenarios Ponderados: Bear / Base / Bull",
        col_widths=[1.8, 1.8, 1.8, 1.8]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 12: ANÁLISIS DE SENSIBILIDAD
    # ═════════════════════════════════════════════
    add_heading_1(doc, "12. Análisis de Sensibilidad Bidimensional y Matriz de Rangos")
    add_heading_2(doc, "12.1 Sensibilidad Cruzada WACC vs. Tasa de Crecimiento Terminal (g)")
    add_body_p(doc, "La matriz bidimensional demuestra que el dictamen COMPRAR se sostiene para cualquier WACC inferior a 8,0% en combinación con tasas de crecimiento terminal g ≥ 1,5%.")
    
    add_paired_figures(
        doc,
        19, "Mapa de Calor WACC vs. g: Matriz de sensibilidad bidimensional del precio por acción en ARS.",
        20, "Simulación Monte Carlo (20.000 iteraciones): Densidad empírica con innovaciones t-Student (ν=4,2)."
    )
    
    add_custom_table(
        doc,
        ["Metodología de Valuación", "Mínimo", "Central", "Máximo"],
        [
            ["DCF λ=1,0 (Riesgo País Completo)", "--", "776", "--"],
            ["DCF Damodaran (g = ROIC × RR)", "--", "993", "--"],
            ["DCF Caso Base Oficial (Dumrauf, λ=0,20)", "--", "1.236", "--"],
            ["DCF Crecimiento Optimista (g=2,5%)", "--", "1.371", "--"],
            ["Monte Carlo (Percentiles P5 / Mediana / P95)", "844", "1.237", "1.737"],
            ["Target Integrado (DCF + Opción Real PEAL V)", "--", "1.255,60", "--"],
            ["Marco de Escenarios Ponderado (25/50/25)", "--", "1.184", "--"],
            ["Cotización Spot de Mercado", "--", "982,50", "--"]
        ],
        caption="Síntesis de Precios Intrínsecos por Metodología (ARS/acción)",
        col_widths=[3.5, 1.1, 1.2, 1.1]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 13: SIMULACIÓN MONTE CARLO Y DCF INVERSO
    # ═════════════════════════════════════════════
    add_heading_1(doc, "13. Simulación Estocástica de Monte Carlo y DCF Inverso")
    add_heading_2(doc, "13.1 Distribución Empírica con Innovaciones t-Student (ν=4,2)")
    add_equation_box(doc, "WACC_{sim} = WACC + σ_{WACC} × √((ν-2)/ν) × T_ν,   g_{sim} = g + σ_g × √((ν-2)/ν) × T_ν", "Innovaciones t-Student")
    add_body_p(doc, "La simulación de Monte Carlo con 20.000 iteraciones e innovaciones con colas pesadas t-Student arroja una mediana de ARS 1.237 y un intervalo de confianza P5–P95 de ARS 844 a ARS 1.737, con un 85,2% de probabilidad empírica de retorno positivo frente al spot.")
    
    add_heading_2(doc, "13.2 Análisis de Expectativas Implícitas: Tasa Implícita del DCF Inverso (g* = 0,69%)")
    add_body_p(doc, "Despejando la tasa de crecimiento terminal implícita en la cotización de mercado, se comprueba que el precio actual descuenta un crecimiento a perpetuidad de apenas g* = 0,69%, muy por debajo del 2,0% fundamental.")
    
    add_paired_figures(
        doc,
        21, "DCF Inverso: Tasa g implícita descontada en el precio spot de mercado (0,69%).",
        22, "Convergencia de Múltiplos: Normalización del múltiplo EV/EBITDA hacia 6,75x."
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 14: VALUACIÓN RELATIVA SECTORIAL
    # ═════════════════════════════════════════════
    add_heading_1(doc, "14. Valuación Relativa Sectorial y Múltiplos Comparables LTM y Forward")
    add_body_p(doc, "El análisis de múltiplos sectoriales confirma la divergencia entre la valuación LTM y la capacidad de generación normalizada. Mientras el múltiplo LTM de Aluar (13,3x) se encuentra transitoriamente afectado por el ciclo de FY2025, su rendimiento del flujo libre de caja proyectado (FCF Yield 2027E > 10%) supera con holgura la mediana internacional (3,8%):")
    
    add_custom_table(
        doc,
        ["Compañía / Ticker", "EV/EBITDA LTM", "Crec. Ingresos LTM", "Margen EBITDA LTM", "ROIC LTM", "FCF Yield 2026E/27E"],
        [
            ["ALUAR (ALUA.BA)", "13,3x", "8,0%", "14,9%", "4,3%", "-4,1% / +12,8%"],
            ["Alcoa Corp. (AA)", "8,5x", "6,5%", "12,0%", "5,5%", "4,8%"],
            ["Chalco (ACH)", "9,0x", "9,2%", "15,1%", "8,0%", "5,2%"],
            ["Kaiser Aluminum (KALU)", "7,2x", "4,1%", "9,8%", "3,5%", "3,0%"],
            ["Norsk Hydro (NHYDY)", "6,4x", "5,8%", "13,5%", "6,2%", "4,1%"],
            ["Constellium (CSTM)", "5,4x", "3,2%", "10,1%", "5,0%", "3,5%"],
            ["Rusal", "5,1x", "-2,0%", "8,4%", "2,1%", "1,5%"],
            ["Mediana Sectorial", "7,2x", "5,0%", "11,1%", "5,3%", "3,8%"]
        ],
        caption="Tabla Comparativa Sectorial y Múltiplos Financieros LTM (julio-2026)",
        col_widths=[2.4, 1.0, 1.0, 1.0, 0.8, 1.0]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 15: GESTIÓN CUANTITATIVA DE RIESGO
    # ═════════════════════════════════════════════
    add_heading_1(doc, "15. Gestión Cuantitativa de Riesgo y Asignación de Portafolio")
    add_heading_2(doc, "15.1 Criterio de Asignación de Kelly y Límite de Tolerancia por CVaR")
    add_equation_box(doc, "f* = (μ - Rf) / σ² = (24,63% - 4,70%) / (52,06%)² = 73,5% (Full-Kelly)   =>   Half-Kelly = 36,8%", "Criterio de Kelly")
    add_body_p(doc, "Para carteras institucionales se impone un límite de política de riesgo del 20,0%, restringido por el Conditional Value-at-Risk mensual al 95% (-32,80%). El Filtro de Kalman dinámico muestra que el Beta reciente converge a 0,764:")
    add_equation_box(doc, "y_t = β_t × x_t + v_t,   β_t = β_{t-1} + w_t,   K_t = P_{t|t-1} × x_t × (x_t × P_{t|t-1} × x_t + R)^{-1}", "Filtro de Kalman")
    
    add_paired_figures(
        doc,
        23, "Criterio de Asignación de Kelly: Full-Kelly (73,5%), Half-Kelly (36,8%) y Límite de Riesgo (20,0%).",
        24, "Beta Dinámico por Filtro de Kalman: Convergencia a 0,764 en el régimen de riesgo reciente."
    )
    
    add_heading_2(doc, "15.2 Comportamiento Histórico de ALUA.BA en Ventanas de Crisis")
    add_custom_table(
        doc,
        ["Episodio de Estrés", "Retorno Acumulado", "Volatilidad Anualizada", "Observaciones", "Comportamiento"],
        [
            ["Crisis Covid-19 (Marzo 2020)", "-44,73%", "112,54%", "50 ruedas", "Shock global de liquidez"],
            ["Volatilidad Cambiaria (Julio 2022)", "+12,44%", "45,91%", "103 ruedas", "Cobertura dolarizada"],
            ["Post-Elecciones 2023", "+18,28%", "105,36%", "59 ruedas", "Revaluación de activos reales"],
            ["Shock LME Post-Pico (2022–2023)", "+13,61%", "44,40%", "200 ruedas", "Resiliencia por costos bajos"]
        ],
        caption="Comportamiento Histórico de ALUA.BA (USD) en Ventanas de Crisis",
        col_widths=[2.2, 1.2, 1.2, 1.1, 1.3]
    )
    
    add_heading_2(doc, "15.3 Evaluación Metodológica de VaR/CVaR y Modelado de Colas (EVT-GPD)")
    add_body_p(doc, "La expansión de Cornish-Fisher subestima el riesgo extremo en la cola del 99% (-13,41%). Por ello se adopta la Teoría de Valores Extremos (EVT) con ajuste Pareto Generalizado (GPD, ξ = 0,214 > 0), capturando la leptocurtosis y colas pesadas de tipo Fréchet:")
    add_equation_box(doc, "G_{ξ, β}(y) = 1 - (1 + ξ × y / β)^{-1/ξ}", "Pareto Generalizado (GPD)")
    
    add_custom_table(
        doc,
        ["Metodología de Riesgo", "VaR 90%", "VaR 95%", "VaR 99%"],
        [
            ["Histórico", "-3,47%", "-4,73%", "-8,20%"],
            ["Normal (paramétrico)", "-4,12%", "-5,31%", "-7,54%"],
            ["t-Student (ν=4,46)", "-3,61%", "-4,99%", "-8,57%"],
            ["Cornish-Fisher", "-2,70%", "-5,22%", "-13,41%"],
            ["CVaR / Expected Shortfall (paramétrico)", "-5,67%", "-6,68%", "-8,65%"],
            ["CVaR Cornish-Fisher (Simulado)", "-7,12%", "-10,49%", "-20,35%"]
        ],
        caption="VaR y CVaR Diario por Metodología (retornos USD, 1 día)",
        col_widths=[3.5, 1.1, 1.1, 1.1]
    )
    
    add_paired_figures(
        doc,
        25, "Teoría de Valores Extremos (EVT-GPD): Estimación de pérdidas de cola con VaR 99% (8,20%) y ES 99% (10,35%).",
        26, "Simulación LME (OU vs. MBG): Proceso de reversión a la media Ornstein-Uhlenbeck (θ = USD 2.814/t)."
    )
    
    add_heading_2(doc, "15.4 Dinámica de Commodities: Ornstein-Uhlenbeck vs. Movimiento Browniano")
    add_equation_box(doc, "dS_t = κ × (θ - S_t) dt + σ × S_t dW_t,   κ = 0,384,   θ = USD 2.814/t", "Ornstein-Uhlenbeck")
    
    add_paired_figures(
        doc,
        27, "DCF Estocástico Continuo: Densidad de probabilidad con 10.000 trayectorias de precios y costos.",
        28, "Matriz Bidimensional de Riesgo: Mapeo según estándar CFA de Probabilidad vs. Impacto."
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 16: MODELIZACIÓN ESTOCÁSTICA Y ECONOMETRÍA
    # ═════════════════════════════════════════════
    add_heading_1(doc, "16. Modelización Estocástica, Dependencia de Cola y Validación Econométrica")
    add_heading_2(doc, "16.1 Validacion de la Curva Soberana: AR(1) vs. Cox-Ingersoll-Ross (CIR)")
    add_equation_box(doc, "dr_t = κ × (θ - r_t) dt + σ × √r_t dW_t,   Condición de Feller: 2κθ > σ²", "Proceso CIR")
    add_body_p(doc, "Con κ = 0,7582, θ = 0,1614 y σ = 0,085, se verifica 2(0,7582)(0,1614) = 0,2448 > 0,0072 = σ².")
    
    add_paired_figures(
        doc,
        29, "Proceso CIR Soberano: Simulación estocástica del EMBI+ respetando la no-negatividad estricta.",
        30, "Beta Condicional GARCH(1,1): Variación temporal de la volatilidad condicional y riesgo sistemático."
    )
    
    add_heading_2(doc, "16.2 Dependencia Asimétrica en Caídas: Cópula de Clayton y Sobol")
    add_equation_box(doc, "C_θ(u, v) = (u^{-θ} + v^{-θ} - 1)^{-1/θ},   λ_L = 2^{-1/θ} = 0,38", "Cópula de Clayton")
    
    add_paired_figures(
        doc,
        31, "Cópula de Clayton Asimétrica: Modelado de co-caídas extremas entre el activo y el mercado (λ_L = 0,38).",
        22, "Convergencia de Múltiplos: Normalización del múltiplo EV/EBITDA hacia 6,75x."
    )
    
    add_heading_2(doc, "16.3 Matriz de Diagnóstico Econométrico y Pruebas Estadísticas")
    add_custom_table(
        doc,
        ["Prueba Estadística", "Variable / Proceso", "Estadístico", "P-Valor / Criterio", "Conclusión Econométrica"],
        [
            ["Dickey-Fuller Aumentado (ADF)", "Retornos Diarios ALUA (USD)", "t = -3,84", "p < 0,01", "Serie estacionaria I(0)"],
            ["Phillips-Perron (PP)", "Retornos Diarios ALUA (USD)", "Z_t = -4,12", "p < 0,01", "Estacionariedad robusta a heterocedasticidad"],
            ["Jarque-Bera", "Normalidad de Retornos", "JB = 4.391,7", "p < 0,0001", "Rechazo de normalidad (Curtosis = 7,82)"],
            ["Kolmogorov-Smirnov", "t-Student (ν=4,46) vs. Normal", "D = 0,018", "p = 0,42", "Ajuste óptimo con colas pesadas t-Student"],
            ["Cópula de Clayton", "Dependencia de Cola Inferior", "ΔAIC = -88,9", "λ_L = 0,38", "Dependencia asimétrica superior a Gaussiana"],
            ["ARCH-LM (Lag 5)", "Heterocedasticidad Condicional", "LM = 48,2", "p < 0,001", "Presencia de volatilidad por conglomerados"],
            ["Ljung-Box Q(10)", "Autocorrelación de Residuos", "Q = 8,45", "p = 0,58", "Ausencia de correlación serial residual"],
            ["Quandt-Andrews", "Ruptura Estructural de Media", "F = 1,14", "p = 0,31", "Parámetros estables en la muestra analizada"]
        ],
        caption="Matriz de Pruebas Estadísticas y Diagnóstico Econométrico",
        col_widths=[1.8, 1.8, 0.9, 1.1, 1.4]
    )
    
    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 17: CONCLUSIÓN
    # ═════════════════════════════════════════════
    add_heading_1(doc, "17. Dictamen Final y Recomendación de Inversión")
    add_body_p(doc, "El dictamen cuantitativo del modelo es COMPRAR, fundamentado en:")
    add_body_p(doc, "• Precio Objetivo Base (DCF Dumrauf): ARS 1.236,00 por acción (USD 0,78, upside de +25,8%).")
    add_body_p(doc, "• Precio Objetivo Integrado (Base + Opción Real PEAL V): ARS 1.255,60 por acción (USD 0,79, upside de +27,8%).")
    add_body_p(doc, "• Rango Probabilístico Monte Carlo (P5–P95): ARS 844 a ARS 1.737 por acción, con un 85,2% de probabilidad empírica de retorno positivo.")
    add_body_p(doc, "• Margen de Seguridad Fundamental: El precio spot descuenta una tasa terminal implícita pesimista (g* = 0,69%), ofreciendo un margen de seguridad de ARS 253,50 por acción frente al valor intrínseco de régimen.")
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════════
    # APÉNDICE: ESTADOS FINANCIEROS AUDITADOS Y PROYECCIONES
    # ═════════════════════════════════════════════
    add_heading_1(doc, "Apéndice: Estados Financieros Auditados y Proyecciones de Flujo de Fondos (USD MM)")
    add_body_p(doc, "Los estados financieros consolidados reproducen las cifras anuales auditadas por Price Waterhouse & Co. S.R.L. convertidas al CCL implícito de cierre de cada ejercicio. El NOPAT proyectado aplica la alícuota estatutaria del 35% (Ley 20.628).")
    
    add_custom_table(
        doc,
        ["Componente Tributario", "Monto (USD MM)", "% sobre EBT"],
        [
            ["Resultado Antes de Impuestos (EBT)", "58,5", "100,0%"],
            ["Impuesto Teórico a Tasa Estatutaria (35%)", "(20,5)", "35,0%"],
            ["Efecto del Ajuste por Inflación Impositivo (NIC 29)", "(23,1)", "39,5%"],
            ["Diferencias de Cambio no Deducibles y Otros", "(5,7)", "9,8%"],
            ["Impuesto a las Ganancias Contabilizado", "(49,3)", "84,3%"]
        ],
        caption="Cuadro de Reconciliación Impositiva FY2025",
        col_widths=[3.5, 1.8, 1.7]
    )
    
    add_custom_table(
        doc,
        ["Línea de Cuenta", "FY2020", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025"],
        [
            ["Ventas Netas (Revenue)", "823,5", "513,5", "654,7", "647,8", "915,9", "1.092,6"],
            ["Costo de Ventas", "(711,2)", "(414,4)", "(448,3)", "(536,0)", "(744,0)", "(930,5)"],
            ["Ganancia Bruta", "112,3", "99,1", "206,4", "111,8", "171,9", "162,1"],
            ["Gastos de Administración y Comercialización", "(63,8)", "(37,8)", "(44,1)", "(50,1)", "(70,1)", "(100,8)"],
            ["Otros Ingresos / Operativos Netos", "(0,4)", "0,0", "(0,2)", "0,4", "50,8", "31,0"],
            ["EBITDA", "116,5", "110,2", "208,8", "103,5", "204,7", "163,3"],
            ["Depreciación y Amortización (D&A)", "(68,4)", "(48,9)", "(46,7)", "(41,5)", "(52,1)", "(71,0)"],
            ["EBIT (Resultado Operativo)", "48,1", "61,3", "162,1", "62,0", "152,6", "92,3"],
            ["Resultado Financiero y Asociadas", "(84,5)", "14,7", "36,2", "88,4", "3,9", "(33,8)"],
            ["Ganancia Antes de Impuestos (EBT)", "(36,4)", "76,0", "198,3", "150,3", "156,5", "58,5"],
            ["Impuesto a las Ganancias", "(7,5)", "(47,9)", "(82,8)", "(11,3)", "(66,7)", "(49,3)"],
            ["Resultado Neto del Ejercicio", "(43,9)", "28,1", "115,5", "139,0", "89,7", "9,2"],
            ["NOPAT (Resultado Operativo Neto de Impuestos)", "31,3", "39,9", "105,4", "40,3", "99,2", "60,0"]
        ],
        caption="Estado de Resultados Consolidado Auditado (FY2020–FY2025, USD MM)",
        col_widths=[2.4, 0.75, 0.75, 0.75, 0.75, 0.75, 0.75]
    )
    
    add_custom_table(
        doc,
        ["Rubro Patrimonial", "FY2020", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025"],
        [
            ["Activo Corriente", "454,6", "318,1", "478,9", "490,1", "902,9", "1.028,1"],
            ["Activo No Corriente (inc. Bienes de Uso)", "597,4", "379,9", "353,6", "422,3", "556,6", "940,9"],
            ["Total Activo", "1.052,0", "698,0", "832,6", "912,4", "1.459,5", "1.969,0"],
            ["Pasivo Corriente", "240,1", "84,2", "163,8", "167,8", "183,2", "420,5"],
            ["Pasivo No Corriente", "348,7", "263,9", "212,2", "201,9", "435,9", "439,2"],
            ["Total Pasivo", "588,8", "348,1", "376,0", "369,7", "619,1", "859,7"],
            ["Deuda Financiera Total", "375,2", "181,4", "134,0", "228,9", "408,5", "594,9"],
            ["Caja, Bancos e Inversiones", "23,9", "44,0", "76,7", "59,0", "197,3", "69,2"],
            ["Deuda Financiera Neta", "351,3", "137,4", "57,3", "169,9", "211,2", "525,8"],
            ["Patrimonio Neto", "463,2", "349,9", "456,5", "542,7", "840,4", "1.109,3"],
            ["Capital Invertido (NOA)", "838,4", "531,3", "590,5", "771,6", "1.248,9", "1.704,3"]
        ],
        caption="Estado de Situación Patrimonial / Balance Consolidado Auditado (FY2020–FY2025, USD MM)",
        col_widths=[2.4, 0.75, 0.75, 0.75, 0.75, 0.75, 0.75]
    )
    
    add_custom_table(
        doc,
        ["Flujo de Caja", "FY2020", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025"],
        [
            ["Flujo de Caja Operativo (FCO)", "245,7", "102,4", "89,2", "113,6", "79,2", "30,7"],
            ["Inversiones en Bienes de Uso (CAPEX)", "(103,0)", "(9,4)", "(11,1)", "(65,9)", "(25,6)", "(243,9)"],
            ["Flujo de Caja de Inversión (FCI)", "(102,1)", "(9,4)", "(11,1)", "(65,9)", "(25,6)", "(244,0)"],
            ["Flujo de Caja de Financiación (FCF)", "(197,8)", "(45,1)", "(34,4)", "(70,8)", "91,1", "20,3"],
            ["Variación Neta de Efectivo", "(54,2)", "47,8", "43,8", "(23,1)", "144,7", "(193,0)"]
        ],
        caption="Estado de Flujos de Efectivo Consolidado Auditado (FY2020–FY2025, USD MM)",
        col_widths=[2.4, 0.75, 0.75, 0.75, 0.75, 0.75, 0.75]
    )
    
    add_custom_table(
        doc,
        ["Cuenta / Año Fiscal", "2026E", "2027E", "2028E", "2029E", "2030E"],
        [
            ["Ventas Netas (Revenue)", "1.591,4", "1.543,0", "1.494,6", "1.446,3", "1.397,9"],
            ["EBITDA", "391,2", "369,3", "347,4", "325,5", "303,6"],
            ["(-) Depreciación y Amortización", "(108,4)", "(105,1)", "(101,8)", "(98,5)", "(95,2)"],
            ["EBIT", "282,8", "264,2", "245,6", "227,0", "208,4"],
            ["(-) Impuestos Operativos (35% Estatutaria)", "(99,0)", "(92,5)", "(86,0)", "(79,4)", "(72,9)"],
            ["NOPAT", "183,8", "171,7", "159,6", "147,6", "135,5"],
            ["(+) D&A (Reversión)", "108,4", "105,1", "101,8", "98,5", "95,2"],
            ["(-) CAPEX de Mantenimiento", "(103,2)", "(100,0)", "(96,9)", "(93,8)", "(90,6)"],
            ["(-) Variación Capital de Trabajo (ΔNWC)", "(246,5)", "23,9", "23,9", "23,9", "23,9"],
            ["FCFF Explícito", "-57,5", "200,7", "188,4", "176,2", "164,0"]
        ],
        caption="Proyecciones del Free Cash Flow to the Firm (FCFF, FY2026E–FY2030E, USD MM)",
        col_widths=[2.8, 0.8, 0.8, 0.8, 0.8, 0.8]
    )
    
    add_body_p(doc, "Reconciliación FCFF Explícito 2030E vs. Terminal Normalizado: FCFF Explícito 2030E: USD 164,0 MM (incluye CAPEX de mantenimiento de USD 90,6 MM y ΔNWC = USD 23,9 MM). FCFF Terminal Normalizado (g = 2,0%): USD 135,5 MM (asume estado estacionario donde CAPEX = D&A y ΔNWC = $0, convergiendo exactamente al NOPAT de 2030E según la metodología Dumrauf Cap. 14).")
    
    # ═════════════════════════════════════════════════════════════════════════
    # REFERENCIAS BIBLIOGRÁFICAS Y DISCLAIMER
    # ═════════════════════════════════════════════
    add_heading_1(doc, "Referencias Bibliográficas y Fuentes Académicas")
    references = [
        "1. Damodaran, Aswath (2012). Investment Valuation: Tools and Techniques for Determining the Value of Any Asset, 3rd Ed., John Wiley & Sons.",
        "2. Dumrauf, Guillermo L. (2013). Finanzas Corporativas: Un Enfoque Latinoamericano, 3ª Ed., Alfaomega Grupo Editor.",
        "3. Ross, Stephen A., Westerfield, Randolph W., Jaffe, Jeffrey (2019). Corporate Finance, 12th Ed., McGraw-Hill Education.",
        "4. Alexander, Gordon J., Sharpe, William F., Bailey, Jeffery V. (2001). Fundamentos de Inversiones, Pearson Educación / Prentice Hall.",
        "5. Black, Fischer, Scholes, Myron (1973). 'The Pricing of Options and Corporate Liabilities', Journal of Political Economy, 81(3), pp. 637–654.",
        "6. Longstaff, Francis A., Schwartz, Eduardo S. (2001). 'Valuing American Options by Simulation: A Simple Least-Squares Approach', Review of Financial Studies, 14(1), pp. 113–147.",
        "7. Blume, Marshall E. (1975). 'Betas and Their Regression Tendencies', Journal of Finance, 30(3), pp. 785–795.",
        "8. Hamada, Robert S. (1972). 'The Effect of the Firm's Capital Structure on the Systematic Risk of Common Stocks', Journal of Finance, 27(2), pp. 435–452.",
        "9. Sklar, Abe (1959). 'Fonctions de Répartition à n Dimensions et Leurs Marges', Publications de l'Institut de Statistique de l'Université de Paris, 8, pp. 229–231.",
        "10. Rockafellar, R. Tyrrell, Uryasev, Stanislav (2000). 'Optimization of Conditional Value-at-Risk', Journal of Risk, 2(3), pp. 21–42.",
        "11. Jaschke, Stefan R. (2001). 'The Cornish-Fisher-Expansion in the Context of Delta-Gamma-Normal Approximations', Journal of Risk, 4(4), pp. 33–52.",
        "12. Engle, Robert F., Granger, Clive W.J. (1987). 'Co-integration and Error Correction: Representation, Estimation and Testing', Econometrica, 55(2), pp. 251–276.",
        "13. Aluar Aluminio Argentino S.A.I.C. (2020–2025). Estados Financieros Consolidados Auditados y Memorias Anuales, Comisión Nacional de Valores (CNV) & Bolsas y Mercados Argentinos (BYMA)."
    ]
    for ref in references:
        add_body_p(doc, ref, space_after=2)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    create_callout_box(
        doc,
        ["Este documento fue elaborado con fines exclusivamente de investigación académica, educativa y de análisis cuantitativo independiente por Federico Agustín Chillón (FCE, Universidad Nacional de Cuyo). El presente reporte no constituye una oferta pública, asesoramiento financiero personalizado, recomendación vinculante de compra o venta, ni intermediación bursátil en los términos de la Ley de Mercado de Capitales N.° 26.831 y las normas reglamentarias de la Comisión Nacional de Valores (CNV). El autor certifica que las opiniones técnicas, estimaciones financieras y modelizaciones cuantitativas reflejan un análisis independiente fundamentado rigurosamente en datos contables públicos auditados por Price Waterhouse & Co. S.R.L., series de mercado de BYMA, CBOE, S&P Dow Jones y el London Metal Exchange (LME). Cualquier decisión de inversión que un tercero adopte en base a este modelo es de su exclusiva responsabilidad y riesgo."],
        title="Aviso Legal y Certificación del Analista (Research Standards / Ley 26.831)"
    )
    
    doc.save(OUTPUT_DOCX)
    print(f"Documento guardado con éxito en: {OUTPUT_DOCX}")
    try:
        shutil.copy2(OUTPUT_DOCX, DOWNLOADS_DOCX)
        print(f"Copia sincronizada en Downloads: {DOWNLOADS_DOCX}")
    except Exception as e:
        print(f"Aviso al copiar en Downloads: {e}")

if __name__ == "__main__":
    build_word_report()
