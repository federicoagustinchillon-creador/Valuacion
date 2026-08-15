# -*- coding: utf-8 -*-
"""
generate_master_word_report.py
Generador Maestro de Máxima Fidelidad (1:1) del Informe de Valuación en Microsoft Word (.docx).
Parsea y replica de forma idéntica la estructura de 16 páginas, tipografía Georgia,
paleta institucional Navy (#0D233A), portada en 2 columnas, índice general en 2 columnas en página 2,
20 secciones, 43 subsecciones, 31 figuras de 300 DPI y 22 tablas estilizadas de
reporte_valuacion.tex / Informe_Valuacion_Aluar_UNCuyo.pdf.
Mapea exactamente la distribución de 16 páginas 1:1 con el PDF oficial de XeLaTeX.
Configura el idioma a español (es-AR) en todos los estilos y nodos XML.
"""

import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import win32com.client

DIR = os.path.dirname(os.path.abspath(__file__))
TEX_PATH = os.path.join(DIR, "tex", "reporte_valuacion.tex")
DOCX_PATH = os.path.join(DIR, "Informe_Valuacion_Aluar_UNCuyo.docx")
FIG_DIR = os.path.join(DIR, "..", "03_Modelo_y_Codigo", "figuras")

# Paleta institucional
HEX_NAVY = "0D233A"
HEX_BLUE = "1E3A8A"
HEX_GREEN = "00843D"
HEX_RED = "B91C1C"
HEX_GOLD = "D97706"
HEX_GRAY = "4A4A4A"
HEX_LIGHT_BG = "F8FAFC"
HEX_BORDER = "D1D5DB"

COLOR_NAVY = RGBColor(13, 35, 58)
COLOR_BLUE = RGBColor(30, 58, 138)
COLOR_GREEN = RGBColor(0, 132, 61)
COLOR_RED = RGBColor(185, 28, 28)
COLOR_GOLD = RGBColor(217, 119, 6)
COLOR_GRAY = RGBColor(74, 74, 74)
COLOR_DARK = RGBColor(33, 37, 41)
COLOR_WHITE = RGBColor(255, 255, 255)
COLOR_BORDER = RGBColor(209, 213, 219)

def set_lang_es(element):
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'es-AR')
    lang.set(qn('w:bidi'), 'ar-SA')
    if hasattr(element, '_r'):
        element._r.get_or_add_rPr().append(lang)
    elif hasattr(element, '_p'):
        element._p.get_or_add_pPr().append(lang)
    elif hasattr(element, '_element'):
        element._element.get_or_add_rPr().append(lang)

def set_cell_background(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=6, bottom=6, left=15, right=15):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, top_color=HEX_NAVY, bottom_color=HEX_NAVY, inside_color=HEX_BORDER):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{top_color}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{bottom_color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{inside_color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_callout_borders(table, left_color=HEX_NAVY, bg_color=HEX_LIGHT_BG):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="20" w:space="0" w:color="{left_color}"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def extract_balanced_braces(text, start_pos):
    idx = start_pos
    while idx < len(text) and text[idx] in ' \t\n':
        idx += 1
    if idx >= len(text) or text[idx] != '{':
        return "", start_pos
    depth = 0
    start = idx + 1
    for i in range(idx, len(text)):
        if text[i] == '{' and (i == 0 or text[i-1] != '\\'):
            depth += 1
        elif text[i] == '}' and (i == 0 or text[i-1] != '\\'):
            depth -= 1
            if depth == 0:
                return text[start:i], i + 1
    return text[start:], len(text)

def replace_macro_1arg(text, macro_name, template_fn):
    target = macro_name + "{"
    while target in text:
        pos = text.find(target)
        brace_pos = pos + len(macro_name)
        arg, end_pos = extract_balanced_braces(text, brace_pos)
        rep = template_fn(arg)
        text = text[:pos] + rep + text[end_pos:]
    return text

def replace_macro_2args(text, macro_name, template_fn):
    target = macro_name + "{"
    while target in text:
        pos = text.find(target)
        brace_pos = pos + len(macro_name)
        arg1, end_arg1 = extract_balanced_braces(text, brace_pos)
        arg2, end_arg2 = extract_balanced_braces(text, end_arg1)
        rep = template_fn(arg1, arg2)
        text = text[:pos] + rep + text[end_arg2:]
    return text

def clean_math_to_unicode(text):
    if not text: return ""
    text = text.replace('$', '')
    delims = [
        (r'\left(', '('), (r'\right)', ')'),
        (r'\left[', '['), (r'\right]', ']'),
        (r'\left\{', '{'), (r'\right\}', '}'),
        (r'\left.', ''), (r'\right.', '')
    ]
    for k, v in delims: text = text.replace(k, v)
    while r'\underbrace{' in text:
        pos = text.find(r'\underbrace{')
        a, end_a = extract_balanced_braces(text, pos + len(r'\underbrace'))
        b = ""
        end_b = end_a
        sub_pos = text[end_a:].find('_{')
        if sub_pos != -1 and sub_pos < 3:
            b_brace_pos = end_a + sub_pos + 1
            b, end_b = extract_balanced_braces(text, b_brace_pos)
        rep = f"{a} [{b}]" if b else a
        text = text[:pos] + rep + text[end_b:]
    text = replace_macro_2args(text, r'\frac', lambda num, den: f"({num} / {den})")
    text = replace_macro_1arg(text, r'\sqrt', lambda rad: f"√({rad})")
    text = replace_macro_1arg(text, r'\mathbf', lambda t: f"**{t}**")
    text = replace_macro_1arg(text, r'\mathit', lambda t: f"*{t}*")
    text = replace_macro_1arg(text, r'\mathrm', lambda t: t)
    text = replace_macro_1arg(text, r'\text', lambda t: t)
    text = replace_macro_1arg(text, r'\hat', lambda t: f"{t}̂")
    text = replace_macro_1arg(text, r'\mathbb', lambda t: t)
    text = replace_macro_1arg(text, r'\mathcal', lambda t: t)
    rep_map = [
        (r'\varepsilon', 'ε'), (r'\epsilon', 'ε'), (r'\lambda', 'λ'), (r'\gamma', 'γ'),
        (r'\theta', 'θ'), (r'\sigma', 'σ'), (r'\delta', 'δ'), (r'\alpha', 'α'),
        (r'\beta', 'β'), (r'\omega', 'ω'), (r'\Omega', 'Ω'), (r'\Sigma', 'Σ'),
        (r'\Delta', 'Δ'), (r'\Pi', 'Π'), (r'\tau', 'τ'), (r'\phi', 'φ'),
        (r'\kappa', 'κ'), (r'\xi', 'ξ'), (r'\zeta', 'ζ'), (r'\eta', 'η'),
        (r'\mu', 'μ'), (r'\nu', 'ν'), (r'\infty', '∞'), (r'\partial', '∂'),
        (r'\implies', '⇒'), (r'\approx', '≈'), (r'\equiv', '≡'), (r'\times', '×'),
        (r'\cdot', '·'), (r'\notin', '∉'), (r'\sim', '≈ '), (r'\mid', ' | '),
        (r'\sum', '∑'), (r'\prod', '∏'), (r'\int', '∫'),
        (r'\le', '≤'), (r'\ge', '≥'), (r'\ne', '≠'), (r'\pm', '±'),
        (r'\to', '→'), (r'\in', '∈'),
        (r'\%', '%'), (r'\&', '&'), (r'\_', '_'), (r'\qquad', '    '), (r'\quad', '  '),
        (r'\,', ' '), (r'\;', ' '), (r'\!', '')
    ]
    for k, v in rep_map: text = text.replace(k, v)
    return text.strip()

def clean_full_latex_text(text):
    if not text: return ""
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    if r'\addcontentsline' in text:
        text = re.sub(r'\\addcontentsline\{[^}]*\}\{[^}]*\}\{[^}]*\}', '', text)
    text = replace_macro_2args(text, r'\texorpdfstring', lambda a, b: a)
    text = replace_macro_1arg(text, r'\textbf', lambda t: f"**{t}**")
    text = replace_macro_1arg(text, r'\textit', lambda t: f"*{t}*")
    text = replace_macro_1arg(text, r'\textsf', lambda t: t)
    text = replace_macro_1arg(text, r'\texttt', lambda t: f"`{t}`")
    text = replace_macro_1arg(text, r'\text', lambda t: t)
    text = replace_macro_1arg(text, r'\footnote', lambda t: f" (Nota: {t})")
    text = replace_macro_1arg(text, r'\cite', lambda t: f"[{t}]")
    text = replace_macro_1arg(text, r'\ref', lambda t: f"[{t}]")
    text = replace_macro_1arg(text, r'\label', lambda t: "")
    text = replace_macro_1arg(text, r'\color', lambda t: "")
    text = replace_macro_1arg(text, r'\vspace', lambda t: "")
    text = replace_macro_1arg(text, r'\hspace', lambda t: "")
    def sub_handler(t):
        sub_map = {'0':'₀', '1':'₁', '2':'₂', '3':'₃', '4':'₄', '5':'₅', '6':'₆', '7':'₇', '8':'₈', '9':'₉'}
        return "".join(sub_map.get(c, c) for c in t)
    text = replace_macro_1arg(text, r'\textsubscript', sub_handler)
    def sup_handler(t):
        if t == 'o': return 'º'
        if t == 'a': return 'ª'
        return t
    text = replace_macro_1arg(text, r'\textsuperscript', sup_handler)
    text = re.sub(r'\\(?:small|scriptsize|footnotesize|normalsize|large|Large|LARGE|huge|Huge|tiny)\b', '', text)
    text = re.sub(r'\\(?:bfseries|sffamily|rmfamily|ttfamily|itshape|centering|raggedleft|raggedright|noindent|par|hfill|bigskip|medskip|smallskip)\b', '', text)
    text = re.sub(r'\\rule\{[^}]*\}\{[^}]*\}', '', text)
    text = re.sub(r'\\titlerule', '', text)
    text = re.sub(r'\\clearpage', '', text)
    text = text.replace(r'\par', '')
    text = text.replace(r'\ ', ' ')
    text = text.replace(r'\0', '0')
    def math_replacer(m): return clean_math_to_unicode(m.group(1))
    text = re.sub(r'\$([^$]+)\$', math_replacer, text)
    text = clean_math_to_unicode(text)
    text = re.sub(r'^\s*\{\s*', '', text)
    text = re.sub(r'\s*\}\s*$', '', text)
    text = text.replace('{', '').replace('}', '')
    text = text.replace('---', '—').replace('--', '–').replace('~', ' ')
    return re.sub(r'[ \t]+', ' ', text).strip()

def add_rich_runs(paragraph, text, font_size=Pt(8.6), color=COLOR_DARK, italic_all=False, bold_all=False):
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for token in tokens:
        if not token: continue
        run = paragraph.add_run()
        run.font.name = 'Georgia'
        run.font.size = font_size
        run.font.color.rgb = color
        set_lang_es(run)
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.font.bold = True
            if italic_all: run.font.italic = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.font.italic = True
            if bold_all: run.font.bold = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size.pt * 0.9)
        else:
            run.text = token
            if bold_all: run.font.bold = True
            if italic_all: run.font.italic = True

def add_heading_1(doc, title):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(5)
    h.paragraph_format.space_after = Pt(1.5)
    h.paragraph_format.keep_with_next = True
    add_rich_runs(h, title, font_size=Pt(10.8), color=COLOR_NAVY, bold_all=True)

def add_heading_2(doc, title):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(3.5)
    h.paragraph_format.space_after = Pt(1)
    h.paragraph_format.keep_with_next = True
    add_rich_runs(h, title, font_size=Pt(9.0), color=COLOR_BLUE, bold_all=True)

def add_heading_3(doc, title):
    h = doc.add_heading(level=3)
    h.paragraph_format.space_before = Pt(2.5)
    h.paragraph_format.space_after = Pt(0.8)
    h.paragraph_format.keep_with_next = True
    add_rich_runs(h, title, font_size=Pt(8.0), color=COLOR_GRAY, bold_all=True)

def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.0)
    p.paragraph_format.line_spacing = 1.03
    add_rich_runs(p, text, font_size=Pt(8.6), color=COLOR_DARK)

def add_equation_paragraph(doc, eq_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    add_rich_runs(p, eq_text, font_size=Pt(8.0), color=COLOR_NAVY, bold_all=True)

def add_list_item(doc, text, item_type='bullet', num_idx=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.0)
    p.paragraph_format.line_spacing = 1.03
    p.paragraph_format.left_indent = Inches(0.16)
    prefix = f"{num_idx}. " if item_type == 'number' else "• "
    r_pre = p.add_run(prefix)
    r_pre.font.name = 'Georgia'
    r_pre.font.bold = True
    r_pre.font.color.rgb = COLOR_NAVY
    set_lang_es(r_pre)
    add_rich_runs(p, text, font_size=Pt(8.0), color=COLOR_DARK)

def add_figure_block(doc, img_name, caption_text, width_inches=4.6):
    img_path = os.path.join(FIG_DIR, img_name)
    if not os.path.exists(img_path): return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(2.5)
    p_img.paragraph_format.space_after = Pt(0.5)
    p_img.paragraph_format.keep_with_next = True
    p_img.add_run().add_picture(img_path, width=Inches(width_inches))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(2.0)
    add_rich_runs(p_cap, caption_text, font_size=Pt(7.0), color=COLOR_GRAY, italic_all=True)

def add_figure_pair_block(doc, img1_name, cap1, img2_name, cap2):
    p1, p2 = os.path.join(FIG_DIR, img1_name), os.path.join(FIG_DIR, img2_name)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in range(2):
        set_cell_margins(tbl.cell(0, c), top=6, bottom=3, left=10, right=10)
        set_cell_margins(tbl.cell(1, c), top=0, bottom=12, left=10, right=10)
    p_img1 = tbl.cell(0, 0).paragraphs[0]
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(p1): p_img1.add_run().add_picture(p1, width=Inches(3.20))
    p_img2 = tbl.cell(0, 1).paragraphs[0]
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(p2): p_img2.add_run().add_picture(p2, width=Inches(3.20))
    p_cap1 = tbl.cell(1, 0).paragraphs[0]
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_runs(p_cap1, cap1, font_size=Pt(6.8), color=COLOR_GRAY, italic_all=True)
    p_cap2 = tbl.cell(1, 1).paragraphs[0]
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_runs(p_cap2, cap2, font_size=Pt(6.8), color=COLOR_GRAY, italic_all=True)

def add_figure_and_callout_pair(doc, img_name, cap_text, box_title, rows_data, box_note):
    p_img_path = os.path.join(FIG_DIR, img_name)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_img, c_box = tbl.cell(0, 0), tbl.cell(0, 1)
    set_cell_margins(c_img, top=8, bottom=8, left=8, right=8)
    set_cell_margins(c_box, top=8, bottom=8, left=8, right=8)
    p_img = c_img.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(p_img_path): p_img.add_run().add_picture(p_img_path, width=Inches(3.20))
    p_cap = c_img.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_runs(p_cap, cap_text, font_size=Pt(6.8), color=COLOR_GRAY, italic_all=True)
    sub_tbl = c_box.add_table(rows=1, cols=1)
    sub_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_callout_borders(sub_tbl, left_color=HEX_NAVY, bg_color=HEX_LIGHT_BG)
    cell_inner = sub_tbl.cell(0, 0)
    set_cell_background(cell_inner, HEX_LIGHT_BG)
    set_cell_margins(cell_inner, top=12, bottom=12, left=15, right=15)
    p_ttl = cell_inner.paragraphs[0]
    p_ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_runs(p_ttl, box_title, font_size=Pt(7.5), color=COLOR_NAVY, bold_all=True)
    if rows_data:
        t_sobol = cell_inner.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        t_sobol.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, r in enumerate(rows_data):
            for c_idx, val in enumerate(r):
                cell_s = t_sobol.cell(r_idx, c_idx)
                set_cell_margins(cell_s, top=5, bottom=5, left=8, right=8)
                if r_idx == 0: set_cell_background(cell_s, HEX_NAVY)
                elif r_idx % 2 == 1: set_cell_background(cell_s, HEX_LIGHT_BG)
                p_s = cell_s.paragraphs[0]
                p_s.paragraph_format.space_after = Pt(0)
                if c_idx > 0: p_s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: p_s.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if r_idx == 0: add_rich_runs(p_s, val, font_size=Pt(6.5), color=COLOR_WHITE, bold_all=True)
                else: add_rich_runs(p_s, val, font_size=Pt(6.5), color=COLOR_DARK)
    if box_note:
        p_nt = c_box.add_paragraph()
        p_nt.paragraph_format.space_before = Pt(1)
        add_rich_runs(p_nt, box_note, font_size=Pt(6.5), color=COLOR_GRAY, italic_all=True)

def add_table_block(doc, caption, rows_data, notes=None, is_appendix=False):
    if not rows_data or len(rows_data) == 0: return
    num_rows = len(rows_data)
    num_cols = max(len(r) for r in rows_data)
    if num_cols == 0: return
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(1.5 if is_appendix else 2.5)
        p_cap.paragraph_format.space_after = Pt(0.5 if is_appendix else 0.8)
        p_cap.paragraph_format.keep_with_next = True
        add_rich_runs(p_cap, f"Tabla: {caption}", font_size=Pt(6.8 if is_appendix else 7.5), color=COLOR_NAVY, bold_all=True)
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    set_table_borders(tbl)
    for r_idx, row in enumerate(rows_data):
        is_header = (r_idx == 0)
        for c_idx in range(num_cols):
            cell = tbl.cell(r_idx, c_idx)
            val = row[c_idx] if c_idx < len(row) else ""
            clean_val = clean_full_latex_text(val)
            set_cell_margins(cell, top=4 if is_appendix else 6, bottom=4 if is_appendix else 6, left=10 if is_appendix else 15, right=10 if is_appendix else 15)
            if is_header: set_cell_background(cell, HEX_NAVY)
            elif r_idx % 2 == 1: set_cell_background(cell, HEX_LIGHT_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            is_numeric = (c_idx > 0) and (any(clean_val.endswith(s) for s in ['%', 'x', '$', 'pb', 'MWh', 't']) or clean_val.replace('.', '').replace(',', '').replace('-', '').replace('+', '').strip().isdigit())
            if is_numeric: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            f_size = Pt(6.2 if is_appendix else 6.8)
            if is_header: add_rich_runs(p, clean_val, font_size=f_size, color=COLOR_WHITE, bold_all=True)
            else: add_rich_runs(p, clean_val, font_size=f_size, color=COLOR_DARK)
    if notes:
        p_nt = doc.add_paragraph()
        p_nt.paragraph_format.space_before = Pt(0.5 if is_appendix else 0.8)
        p_nt.paragraph_format.space_after = Pt(1.0 if is_appendix else 2.0)
        add_rich_runs(p_nt, notes, font_size=Pt(6.0 if is_appendix else 6.5), color=COLOR_GRAY, italic_all=True)

def add_callout_box(doc, title, paragraphs, border_color=HEX_NAVY, bg_color=HEX_LIGHT_BG):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_callout_borders(tbl, left_color=border_color, bg_color=bg_color)
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=14, bottom=14, left=25, right=25)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0.8)
    add_rich_runs(p0, title, font_size=Pt(7.8), color=COLOR_NAVY, bold_all=True)
    for p_text in paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0.8)
        p.paragraph_format.line_spacing = 1.02
        add_rich_runs(p, p_text, font_size=Pt(7.4), color=COLOR_DARK)

def configure_section_headers_footers(section, is_first=False):
    section.top_margin = Inches(0.63)
    section.bottom_margin = Inches(0.63)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(0.63)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    
    if not is_first:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.text = ""
        hrun = hp.add_run("ALUAR S.A.I.C. (ALUA.BA) · Equity Research\t\tFCE · UNCuyo")
        hrun.font.name = "Georgia"
        hrun.font.size = Pt(8.0)
        hrun.font.bold = True
        hrun.font.color.rgb = COLOR_NAVY
        set_lang_es(hrun)
        
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.text = ""
        frun = fp.add_run("Federico Agustín Chillón | federico.chillon@fce.uncu.edu.ar\t\tPágina ")
        frun.font.name = "Georgia"
        frun.font.size = Pt(7.5)
        frun.font.color.rgb = COLOR_GRAY
        set_lang_es(frun)
        
        fldPage = OxmlElement('w:fldSimple')
        fldPage.set(qn('w:instr'), 'PAGE')
        fp._p.append(fldPage)

# Mapeo de puntos de inicio de página 1:1 con LaTeX (16 páginas exactas)
PAGE_BREAK_TRIGGERS = [
    r'\subsection{Análisis Estratégico FODA',                                                    # -> Página 4
    r'\subsection{Mecánica Operativa de Ingresos',                                              # -> Página 5
    r'\subsection{Diagnóstico Econométrico: Test de Cointegración de Engle-Granger',            # -> Página 6
    r'\subsection{Impacto del Ajuste de Carbono en Frontera',                                   # -> Página 7
    r'\section{Costo Promedio Ponderado del Capital (WACC)',                                     # -> Página 8
    r'\section{Puente de Valuación',                                                            # -> Página 9
    r'\section{Valuación por Flujo de Fondos Descontados (DCF)',                                 # -> Página 10
    r'\section{Simulación Estocástica de Monte Carlo y DCF Inverso',                            # -> Página 11
    r'\section{Gestión de Riesgos de Mercado y Asignación Óptima de Cartera',                  # -> Página 12
    r'\section{Modelización Estocástica Avanzada y Validación Econométrica',                    # -> Página 13
    r'\subsection{Dependencia Asimétrica en Caídas: Selección de Cópulas y Sensibilidad de Sobol', # -> Página 14
    r'\section*{Apéndice: Estados Financieros Auditados',                                       # -> Página 15
    r'\section*{Referencias Bibliográficas y Fuentes de Información'                             # -> Página 16
]

def should_break_page_before(line):
    for trigger in PAGE_BREAK_TRIGGERS:
        if trigger in line:
            return True
    return False

def build_full_word_document():
    print("Iniciando construcción 1:1 de Informe_Valuacion_Aluar_UNCuyo.docx...")
    doc = docx.Document()
    
    # 0. Defaults de idioma es-AR y Metadatos Oficiales (Macrodatos)
    cp = doc.core_properties
    cp.author = "Federico Agustín Chillón"
    cp.title = "Aluar Aluminio Argentino - Informe de Valuación y Análisis de Riesgo"
    cp.subject = "Valuación Fundamental por Flujo de Fondos Descontados y Modelos Estocásticos (FCE - UNCuyo)"
    cp.keywords = "Valuación, Aluar, ALUA.BA, DCF, WACC, Opciones Reales, Monte Carlo, GARCH, Cópula Clayton, FCE UNCuyo, Chillón"
    cp.comments = "Informe de Valuación Financiera y Análisis Cuantitativo de Riesgo elaborado por Federico Agustín Chillón para la Facultad de Ciencias Económicas de la Universidad Nacional de Cuyo."
    cp.category = "Equity Research / Finanzas Corporativas"
    cp.language = "es-AR"
    cp.last_modified_by = "Federico Agustín Chillón"
    
    try:
        rPrDefault = doc.styles.element.xpath('.//w:docDefaults/w:rPrDefault/w:rPr')
        if rPrDefault:
            rPrDefault[0].append(parse_xml(f'<w:lang {nsdecls("w")} w:val="es-AR" w:bidi="ar-SA" w:eastAsia="es-AR"/>'))
        for s in doc.styles:
            if hasattr(s, 'element') and hasattr(s.element, 'rPr'):
                s.element.get_or_add_rPr().append(parse_xml(f'<w:lang {nsdecls("w")} w:val="es-AR" w:bidi="ar-SA" w:eastAsia="es-AR"/>'))
    except Exception as e:
        print(f"Nota docDefaults: {e}")
        
    # Sección 1: Portada (1 col, sin encabezado)
    sec_cover = doc.sections[0]
    sec_cover.different_first_page_header_footer = True
    configure_section_headers_footers(sec_cover, is_first=True)
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Georgia'
    style_normal.font.size = Pt(8.6)
    style_normal.font.color.rgb = COLOR_DARK
    
    # ==========================================
    # HOJA 1: PORTADA 1:1 EXACTA
    # ==========================================
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l, cell_r = hdr_table.cell(0, 0), hdr_table.cell(0, 1)
    set_cell_margins(cell_l, 0, 0, 0, 0)
    set_cell_margins(cell_r, 0, 0, 0, 0)
    
    p_tl = cell_l.paragraphs[0]
    p_tl.paragraph_format.space_after = Pt(1)
    r_t1 = p_tl.add_run("Aluar Aluminio Argentino\n")
    r_t1.font.name = "Georgia"
    r_t1.font.size = Pt(18)
    r_t1.font.bold = True
    r_t1.font.color.rgb = COLOR_NAVY
    set_lang_es(r_t1)
    
    r_t2 = p_tl.add_run("Informe de Valuación y Análisis de Riesgo")
    r_t2.font.name = "Georgia"
    r_t2.font.size = Pt(10.5)
    r_t2.font.color.rgb = COLOR_GRAY
    set_lang_es(r_t2)
    
    p_tr = cell_r.paragraphs[0]
    p_tr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_tr.paragraph_format.space_after = Pt(1)
    r_r1 = p_tr.add_run("Análisis Financiero\nFCE · UNCuyo\n5 de agosto de 2026")
    r_r1.font.name = "Georgia"
    r_r1.font.size = Pt(8.5)
    r_r1.font.color.rgb = COLOR_GRAY
    set_lang_es(r_r1)
    
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(2)
    p_div.paragraph_format.space_after = Pt(2)
    r_div = p_div.add_run("―" * 68)
    r_div.font.size = Pt(9)
    r_div.font.bold = True
    r_div.font.color.rgb = COLOR_NAVY
    
    p_aviso = doc.add_paragraph()
    p_aviso.paragraph_format.space_after = Pt(4)
    p_aviso.paragraph_format.line_spacing = 1.02
    add_rich_runs(p_aviso, "**Aviso legal inicial:** Documento de análisis financiero y estudio académico independiente. No constituye oferta de compra, venta ni asesoramiento financiero en los términos de la Ley N.º 26.831. Las valuaciones reflejan estimaciones del autor fundamentadas en estados contables públicos y modelos financieros estocásticos. Ver el descargo legal completo y certificación del analista en la sección final de este informe.", font_size=Pt(7.0), color=COLOR_GRAY, italic_all=True)
    
    # 2 Columnas de Portada
    cover_table = doc.add_table(rows=1, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left, c_right = cover_table.cell(0, 0), cover_table.cell(0, 1)
    set_cell_margins(c_left, 0, 0, 0, 40)
    set_cell_margins(c_right, 0, 0, 40, 0)
    
    # Columna Izquierda: Resumen Ejecutivo y Tesis
    p_lh1 = c_left.paragraphs[0]
    p_lh1.paragraph_format.space_before = Pt(0)
    p_lh1.paragraph_format.space_after = Pt(1.5)
    add_rich_runs(p_lh1, "Resumen Ejecutivo", font_size=Pt(10), color=COLOR_NAVY, bold_all=True)
    
    p_lt1 = c_left.add_paragraph()
    p_lt1.paragraph_format.space_after = Pt(4)
    p_lt1.paragraph_format.line_spacing = 1.05
    add_rich_runs(p_lt1, "El modelo teórico de valuación aplicado a Aluar Aluminio Argentino S.A.I.C. (ALUA.BA) concluye un dictamen teórico de **COMPRAR** con un precio objetivo teórico a 12 meses de **ARS 1.236,00** por acción (Target Integrado **ARS 1.255,60** al incorporar la opción real de expansión eólica PEAL V), frente a una cotización de mercado de ARS 982,50 (+25,8% / +27,8% de retorno esperado). El precio actual descuenta una tasa terminal implícita excesivamente contractiva (g* = 0,69%), subestimando la ventaja en costos de energía autogenerada y el proceso de desapalancamiento operativo proyectado.", font_size=Pt(8.2), color=COLOR_DARK)
    
    p_lh2 = c_left.add_paragraph()
    p_lh2.paragraph_format.space_before = Pt(2)
    p_lh2.paragraph_format.space_after = Pt(1.5)
    add_rich_runs(p_lh2, "Fundamentos Centrales de la Tesis", font_size=Pt(10), color=COLOR_NAVY, bold_all=True)
    
    thesis_points = [
        "**Estructura de Costos en Primer Cuartil Global:** Matriz energética 96% autogenerada (78% renovable vía Futaleufú y PEAL), con costo en efectivo C1 en USD 1.680/t (percentil 29% mundial).",
        "**Ingresos en Moneda Dura con Costos Fijos en Pesos:** 80% de las ventas se exportan a precios fijados en el LME, mientras que los costos operativos locales se licúan en términos reales ante depreciaciones del tipo de cambio.",
        "**Compresión del Diferencial Soberano y Desapalancamiento:** La caída del spread EMBI+ (441 pb) sitúa el WACC en 7,06% (λ = 0,20), y el cese de inversiones intensivas reduce la deuda neta a <1,0x EBITDA hacia FY2028E."
    ]
    for idx, tp in enumerate(thesis_points, 1):
        p_tpi = c_left.add_paragraph()
        p_tpi.paragraph_format.space_after = Pt(2)
        p_tpi.paragraph_format.line_spacing = 1.04
        p_tpi.paragraph_format.left_indent = Inches(0.15)
        r_num = p_tpi.add_run(f"{idx}. ")
        r_num.font.name = "Georgia"
        r_num.font.bold = True
        r_num.font.color.rgb = COLOR_NAVY
        set_lang_es(r_num)
        add_rich_runs(p_tpi, tp, font_size=Pt(8.0), color=COLOR_DARK)
        
    # Columna Derecha: Recuadro Dictamen y Tabla de Parámetros Clave
    box_tbl = c_right.add_table(rows=1, cols=1)
    box_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_callout_borders(box_tbl, left_color=HEX_NAVY, bg_color=HEX_LIGHT_BG)
    box_cell = box_tbl.cell(0, 0)
    set_cell_background(box_cell, HEX_LIGHT_BG)
    set_cell_margins(box_cell, top=30, bottom=30, left=40, right=40)
    
    p_bx0 = box_cell.paragraphs[0]
    p_bx0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bx0.paragraph_format.space_after = Pt(0)
    add_rich_runs(p_bx0, "Dictamen del Modelo Teórico", font_size=Pt(7.5), color=COLOR_GRAY)
    
    p_bx1 = box_cell.add_paragraph()
    p_bx1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bx1.paragraph_format.space_after = Pt(1)
    add_rich_runs(p_bx1, "COMPRAR", font_size=Pt(14), color=COLOR_GREEN, bold_all=True)
    
    p_bx2 = box_cell.add_paragraph()
    p_bx2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bx2.paragraph_format.space_after = Pt(0)
    add_rich_runs(p_bx2, "Target Base (Dumrauf): **ARS 1.236,00** (USD 0,78)", font_size=Pt(7.5), color=COLOR_DARK)
    
    p_bx3 = box_cell.add_paragraph()
    p_bx3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bx3.paragraph_format.space_after = Pt(0)
    add_rich_runs(p_bx3, "Opción Real PEAL V: **+ARS 19,60** (USD 0,01)", font_size=Pt(7.5), color=COLOR_DARK)
    
    p_bx4 = box_cell.add_paragraph()
    p_bx4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bx4.paragraph_format.space_after = Pt(0)
    add_rich_runs(p_bx4, "Target Integrado: **ARS 1.255,60** (USD 0,79)", font_size=Pt(8.2), color=COLOR_NAVY, bold_all=True)
    
    p_bx5 = box_cell.add_paragraph()
    p_bx5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bx5.paragraph_format.space_after = Pt(2)
    add_rich_runs(p_bx5, "Cotización de Mercado: ARS 982,50 | **Retorno: +27,8%**", font_size=Pt(7.5), color=COLOR_DARK)
    
    p_sep = box_cell.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(1)
    p_sep.paragraph_format.space_after = Pt(1)
    r_sp = p_sep.add_run("―" * 32)
    r_sp.font.size = Pt(7)
    r_sp.font.color.rgb = COLOR_NAVY
    
    params_data = [
        ["Costo de Capital (WACC):", "7,06%"],
        ["Costo del Equity (Ke, λ = 0,20):", "9,30%"],
        ["Beta Apalancado (Hamada):", "0,888"],
        ["Prima de Riesgo de Mercado (ERP):", "4,18%"],
        ["Prima País Ponderada (λ × EMBI+):", "0,88%"],
        ["Costo de la Deuda Post-Impuestos (Kd):", "2,47%"],
        ["Estructura D/E Objetivo:", "0,486"],
        ["Crecimiento Terminal (g):", "2,00%"],
        ["Potencial Base / Integrado:", "+25,8% / +27,8%"],
        ["Tipo de Cambio Implícito (CCL):", "1.584,25"],
    ]
    t_params = box_cell.add_table(rows=len(params_data), cols=2)
    t_params.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (k, v) in enumerate(params_data):
        c0, c1 = t_params.cell(r_idx, 0), t_params.cell(r_idx, 1)
        set_cell_margins(c0, 8, 8, 10, 10)
        set_cell_margins(c1, 8, 8, 10, 10)
        p0, p1 = c0.paragraphs[0], c1.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_rich_runs(p0, k, font_size=Pt(7.0), color=COLOR_DARK, bold_all=True)
        add_rich_runs(p1, v, font_size=Pt(7.0), color=COLOR_NAVY, bold_all=True)
        
    p_sep2 = box_cell.add_paragraph()
    p_sep2.paragraph_format.space_before = Pt(1)
    p_sep2.paragraph_format.space_after = Pt(1)
    r_sp2 = p_sep2.add_run("―" * 32)
    r_sp2.font.size = Pt(7)
    r_sp2.font.color.rgb = COLOR_BORDER
    
    p_cal = box_cell.add_paragraph()
    p_cal.paragraph_format.space_after = Pt(0)
    p_cal.paragraph_format.line_spacing = 1.02
    add_rich_runs(p_cal, "**Nota de calibración.** Valuación fundamental según metodología de Dumrauf (Cap. 14) con CRP = λ × EMBI+ y Opción Real por mínimos cuadrados de Longstaff-Schwartz. Datos al 23-jul-2026.", font_size=Pt(6.5), color=COLOR_GRAY, italic_all=True)
    
    # Cuadro Inferior de Portada: Divergencias Clave
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    divergencias_bullets = [
        "**Normalización del EBITDA vs. Múltiplo Deprimido:** El mercado asigna un múltiplo histórico deprimido (13,4x EV/EBITDA) sobre el ejercicio recesivo FY2025; nuestro modelo demuestra que la normalización operativa en FY2026E (EBITDA USD 391,2 MM) reduce el múltiplo forward a 5,2x.",
        "**Desapalancamiento Acelerado Post-CAPEX:** Tras finalizar la etapa intensiva del parque eólico, el flujo de caja operativo reduce la Deuda Neta a <1,0x EBITDA hacia FY2028E, liberando recursos para dividendos.",
        "**Cobertura Operativa y Factor λ = 0,20:** El mercado sobreestima la prima por riesgo soberano al asumir λ = 1,0; el 80% de ventas en moneda dura y costos fijos en pesos justifican un costo de capital de 7,06%."
    ]
    add_callout_box(doc, "Divergencias Clave frente a las Expectativas de Mercado (Tesis de Diferenciación)", divergencias_bullets, border_color=HEX_BLUE, bg_color=HEX_LIGHT_BG)
    
    # ==========================================
    # HOJA 2: ÍNDICE GENERAL EN 2 COLUMNAS (SECCIÓN 2)
    # ==========================================
    sec_toc = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_headers_footers(sec_toc)
    
    c_node = OxmlElement('w:cols')
    c_node.set(qn('w:num'), '2')
    c_node.set(qn('w:space'), '400')
    sec_toc._sectPr.append(c_node)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    add_rich_runs(p_title, "Índice General de Secciones y Contenidos", font_size=Pt(12.0), color=COLOR_NAVY, bold_all=True)
    
    p_r1 = doc.add_paragraph()
    p_r1.paragraph_format.space_after = Pt(3)
    r1 = p_r1.add_run("―" * 32)
    r1.font.size = Pt(8)
    r1.font.color.rgb = COLOR_NAVY
    
    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.space_before = Pt(2)
    p_toc.paragraph_format.space_after = Pt(3)
    run = p_toc.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
    run._r.append(fldSimple)
    
    p_r2 = doc.add_paragraph()
    p_r2.paragraph_format.space_before = Pt(3)
    p_r2.paragraph_format.space_after = Pt(2)
    r2 = p_r2.add_run("―" * 32)
    r2.font.size = Pt(8)
    r2.font.color.rgb = COLOR_GRAY
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(0)
    add_rich_runs(p_desc, "**Estructura del Documento:** Comprende 17 Secciones de Análisis, Apéndice Contable con Estados Financieros Auditados por PwC (FY2020–FY2025), Proyecciones Explícitas de Flujo de Fondos (FCFF 2026E–2030E), 31 Figuras Analíticas con Narrativa Contextual, Referencias Bibliográficas y Sección Separada de Aviso Legal y Certificación del Analista.", font_size=Pt(7.2), color=COLOR_GRAY)
    
    # ==========================================
    # SECCIÓN 3: CUERPO DEL INFORME (PÁGINAS 3 A 16, 1 COLUMNA)
    # ==========================================
    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_headers_footers(sec_body)
    cols_body = sec_body._sectPr.xpath('.//w:cols')
    if cols_body:
        cols_body[0].set(qn('w:num'), '1')
    else:
        c_node1 = OxmlElement('w:cols')
        c_node1.set(qn('w:num'), '1')
        sec_body._sectPr.append(c_node1)
        
    with open(TEX_PATH, 'r', encoding='utf-8') as f:
        tex_full = f.read()
        
    body_match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', tex_full, re.DOTALL)
    tex_body = body_match.group(1) if body_match else tex_full
    
    s1_pos = tex_body.find(r'\section{Descripción de la')
    if s1_pos == -1: s1_pos = tex_body.find(r'\section{Descripci')
    if s1_pos != -1: tex_body = tex_body[s1_pos:]
    
    lines = tex_body.splitlines()
    i = 0
    in_table, table_caption, table_notes, table_lines = False, "", "", []
    in_figure, figure_lines = False, []
    in_tcolorbox, tcolorbox_title, tcolorbox_lines = False, "", []
    in_itemize, in_enumerate, enum_idx = False, False, 1
    fig_counter = 1
    is_appendix_active = False
    
    while i < len(lines):
        line = lines[i].strip()
        if not line or (line.startswith('%') and not line.startswith('% Tabla')):
            i += 1
            continue
            
        if line.startswith(r'\clearpage'):
            i += 1
            continue
            
        # Verificar si esta línea inicia una nueva página según el presupuesto de 16 páginas 1:1
        if should_break_page_before(line):
            doc.add_page_break()
            
        if r'\begin{table}' in line or r'\begin{table*}' in line:
            in_table, table_caption, table_notes, table_lines = True, "", "", []
            i += 1
            continue
            
        if in_table:
            if r'\caption' in line:
                m_cap = re.search(r'\\caption\{([^}]+)\}', line)
                if m_cap: table_caption = clean_full_latex_text(m_cap.group(1))
            elif r'\end{table}' in line or r'\end{table*}' in line:
                in_table = False
                tab_full_str = " ".join(table_lines)
                m_tab = re.search(r'\\begin\{tabular(?:x)?\}(?:\{[^}]*\})?(?:\{[^}]*\})?(.*?)\\end\{tabular(?:x)?\}', tab_full_str, re.DOTALL)
                if m_tab:
                    raw_content = m_tab.group(1)
                    if r'\toprule' in raw_content: raw_content = raw_content.split(r'\toprule', 1)[1]
                    elif r'\hline' in raw_content: raw_content = raw_content.split(r'\hline', 1)[1]
                    raw_rows = [r.strip() for r in raw_content.split(r'\\') if r.strip()]
                    rows_data = []
                    for r in raw_rows:
                        r_clean = re.sub(r'\\(?:toprule|midrule|bottomrule|hline|rowcolor\{[^}]*\}|multicolumn\{[^}]*\}\{[^}]*\})', '', r).strip()
                        if r_clean and '&' in r_clean:
                            cells = [clean_full_latex_text(c) for c in r_clean.split('&')]
                            if len(cells) > 0 and any(len(c) > 0 for c in cells):
                                rows_data.append(cells)
                    if rows_data:
                        add_table_block(doc, table_caption, rows_data, notes=table_notes if table_notes else None, is_appendix=is_appendix_active)
            else:
                if 'Fuente:' in line or 'Nota:' in line or r'\footnotesize' in line or r'\scriptsize' in line or r'\tiny' in line:
                    table_notes = clean_full_latex_text(line)
                table_lines.append(line)
            i += 1
            continue
            
        if r'\begin{figure}' in line or r'\begin{figure*}' in line:
            in_figure, figure_lines = True, []
            i += 1
            continue
            
        if in_figure:
            if r'\end{figure}' in line or r'\end{figure*}' in line:
                in_figure = False
                fig_full_str = "\n".join(figure_lines)
                mini_matches = re.findall(r'\\begin\{minipage\}[^\n]*\n(.*?)\\end\{minipage\}', fig_full_str, re.DOTALL)
                if len(mini_matches) >= 2:
                    if r'\includegraphics' in mini_matches[0] and r'\includegraphics' in mini_matches[1]:
                        imgs, caps = [], []
                        for mm in mini_matches[:2]:
                            m_img = re.search(r'\\includegraphics\[[^\]]*\]\{([^}]+)\}', mm)
                            m_c = re.search(r'\\caption\{([^}]+)\}', mm)
                            if m_img: imgs.append(os.path.basename(m_img.group(1)))
                            if m_c: caps.append(clean_full_latex_text(m_c.group(1)))
                            else: caps.append(f"Figura {fig_counter}")
                        if len(imgs) == 2:
                            add_figure_pair_block(doc, imgs[0], caps[0], imgs[1], caps[1])
                            fig_counter += 2
                    elif r'\includegraphics' in mini_matches[0] and (r'\begin{tcolorbox}' in mini_matches[1] or 'Sobol' in mini_matches[1]):
                        m_img = re.search(r'\\includegraphics\[[^\]]*\]\{([^}]+)\}', mini_matches[0])
                        m_c = re.search(r'\\caption\{([^}]+)\}', mini_matches[0])
                        img_name = os.path.basename(m_img.group(1)) if m_img else "figura_31.png"
                        cap_text = clean_full_latex_text(m_c.group(1)) if m_c else "Figura 31: Cópula de Clayton Asimétrica."
                        sobol_rows = [
                            ["Variable", "Índice Si", "Total STi"],
                            ["Precio Aluminio LME", "0,421", "0,510"],
                            ["Spread Soberano (EMBI)", "0,285", "0,392"],
                            ["Costo Energía / Alúmina", "0,112", "0,165"],
                            ["Tipo de Cambio Real", "0,084", "0,140"]
                        ]
                        sobol_note = "Los índices de Sobol confirman que el 70,6% del Enterprise Value depende conjuntamente del precio del LME y del spread soberano."
                        add_figure_and_callout_pair(doc, img_name, cap_text, "Descomposición de Varianza (Sobol)", sobol_rows, sobol_note)
                        fig_counter += 1
                else:
                    m_img = re.search(r'\\includegraphics\[[^\]]*\]\{([^}]+)\}', fig_full_str)
                    m_c = re.search(r'\\caption\{([^}]+)\}', fig_full_str)
                    if m_img:
                        img_name = os.path.basename(m_img.group(1))
                        cap_text = clean_full_latex_text(m_c.group(1)) if m_c else f"Figura {fig_counter}: Análisis cuantitativo de Aluar."
                        add_figure_block(doc, img_name, cap_text, width_inches=4.6)
                        fig_counter += 1
            else:
                figure_lines.append(line)
            i += 1
            continue
            
        if r'\begin{tcolorbox}' in line:
            in_tcolorbox, tcolorbox_title, tcolorbox_lines = True, "", []
            m_ttl = re.search(r'title=\{([^}]+)\}', line)
            if m_ttl: tcolorbox_title = clean_full_latex_text(m_ttl.group(1))
            i += 1
            continue
            
        if in_tcolorbox:
            if r'\end{tcolorbox}' in line:
                in_tcolorbox = False
                clean_tcb_pars = []
                for tl in tcolorbox_lines:
                    cl = clean_full_latex_text(tl)
                    if cl and len(cl) > 2 and not cl.startswith(r'\begin') and not cl.startswith(r'\end'):
                        clean_tcb_pars.append(cl)
                if not tcolorbox_title: tcolorbox_title = "Dictamen del Modelo Teórico y Conclusiones"
                add_callout_box(doc, tcolorbox_title, clean_tcb_pars, border_color=HEX_NAVY, bg_color=HEX_LIGHT_BG)
            else:
                tcolorbox_lines.append(line)
            i += 1
            continue
            
        if line.startswith(r'\section'):
            brace_pos = line.find('{')
            if brace_pos != -1:
                title, _ = extract_balanced_braces(line, brace_pos)
                clean_ttl = clean_full_latex_text(title)
                if 'Apéndice' in clean_ttl or 'Apendice' in clean_ttl:
                    is_appendix_active = True
                elif 'Referencias' in clean_ttl:
                    is_appendix_active = False
                add_heading_1(doc, clean_ttl)
            i += 1
            continue
            
        if line.startswith(r'\subsection'):
            brace_pos = line.find('{')
            if brace_pos != -1:
                title, _ = extract_balanced_braces(line, brace_pos)
                add_heading_2(doc, clean_full_latex_text(title))
            i += 1
            continue
            
        if line.startswith(r'\subsubsection'):
            brace_pos = line.find('{')
            if brace_pos != -1:
                title, _ = extract_balanced_braces(line, brace_pos)
                add_heading_3(doc, clean_full_latex_text(title))
            i += 1
            continue
            
        if r'\begin{enumerate}' in line:
            in_enumerate, enum_idx = True, 1
            i += 1
            continue
        if r'\end{enumerate}' in line:
            in_enumerate = False
            i += 1
            continue
        if r'\begin{itemize}' in line:
            in_itemize = True
            i += 1
            continue
        if r'\end{itemize}' in line:
            in_itemize = False
            i += 1
            continue
            
        if line.startswith(r'\item'):
            item_txt = clean_full_latex_text(re.sub(r'^\\item\s*', '', line))
            if in_enumerate:
                add_list_item(doc, item_txt, item_type='number', num_idx=enum_idx)
                enum_idx += 1
            else:
                add_list_item(doc, item_txt, item_type='bullet')
            i += 1
            continue
            
        if line.startswith(r'\[') or line.startswith(r'\begin{equation}'):
            eq_lines = [line]
            while i + 1 < len(lines) and not (r'\]' in lines[i] or r'\end{equation}' in lines[i]):
                i += 1
                eq_lines.append(lines[i].strip())
            eq_str = " ".join(eq_lines)
            eq_str = re.sub(r'\\(?:begin\{equation\}|end\{equation\}|\[|\])', '', eq_str)
            clean_eq = clean_full_latex_text(eq_str)
            if clean_eq: add_equation_paragraph(doc, clean_eq)
            i += 1
            continue
            
        if not line.startswith(r'\begin') and not line.startswith(r'\end') and not line.startswith(r'\caption') and not line.startswith(r'\centering') and not line.startswith(r'\vspace') and not line.startswith(r'\noindent') and not line.startswith(r'\hfill') and not line.startswith(r'\includegraphics') and not line.startswith(r'\addcontentsline'):
            clean_p = clean_full_latex_text(line)
            if clean_p and len(clean_p) > 3:
                add_body_paragraph(doc, clean_p)
        i += 1
        
    doc.save(DOCX_PATH)
    size_mb = os.path.getsize(DOCX_PATH) / 1e6
    print(f"[EXITO] Documento Word generado en {DOCX_PATH} ({size_mb:.2f} MB).")
    
    # Actualizar Tabla de Contenidos mediante Word COM Automation
    try:
        print("Actualizando Tabla de Contenidos y campos mediante Word COM Automation...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = os.path.abspath(DOCX_PATH)
        wdoc = word.Documents.Open(abs_path)
        
        # Ajustar estilos TOC 1 y TOC 2 en Word COM (constantes universales -20 y -21)
        try:
            s1 = wdoc.Styles(-20) # wdStyleTOC1
            s1.Font.Name = "Georgia"
            s1.Font.Size = 7.0
            s1.ParagraphFormat.SpaceBefore = 0.5
            s1.ParagraphFormat.SpaceAfter = 0.5
            s1.ParagraphFormat.LineSpacingRule = 0 # wdLineSpaceSingle
            
            s2 = wdoc.Styles(-21) # wdStyleTOC2
            s2.Font.Name = "Georgia"
            s2.Font.Size = 6.0
            s2.ParagraphFormat.SpaceBefore = 0
            s2.ParagraphFormat.SpaceAfter = 0
            s2.ParagraphFormat.LineSpacingRule = 0 # wdLineSpaceSingle
        except Exception as se:
            print(f"Nota ajuste estilo COM: {se}")
            
        for toc in wdoc.TablesOfContents:
            toc.Update()
        wdoc.Fields.Update()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
        print("[EXITO] Tabla de Contenidos y campos de Word actualizados perfectamente vía COM.")
    except Exception as e:
        print(f"Nota sobre Word COM: {e}")

if __name__ == '__main__':
    build_full_word_document()
